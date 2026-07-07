"""Verify that final clean artifacts contain the latest experiment metrics."""
from __future__ import annotations

import json
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def docx_text(path: Path) -> str:
    doc = Document(path)
    chunks = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(chunks)


def pptx_text(path: Path) -> str:
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    chunks: list[str] = []
    with ZipFile(path) as archive:
        names = [
            name
            for name in archive.namelist()
            if (name.startswith("ppt/slides/slide") or name.startswith("ppt/notesSlides/notesSlide"))
            and name.endswith(".xml")
        ]
        for name in sorted(names):
            root = ET.fromstring(archive.read(name))
            chunks.extend(node.text for node in root.findall(".//a:t", ns) if node.text)
    return "\n".join(chunks)


def expected_tokens(metrics: dict, summary: dict) -> list[tuple[str, str]]:
    primary = metrics["models"][metrics["primary_model"]]
    baseline = metrics["models"]["baseline_kitti_logreg"]
    ablation = metrics["models"]["ablation_without_3d_geometry"]
    return [
        ("primary_model", metrics["primary_model"]),
        ("dataset_total", str(summary["num_scenes"])),
        ("positive_critical_scenes", str(summary["positive_critical_scenes"])),
        ("negative_scenes", str(summary["negative_scenes"])),
        ("train_split", str(summary["split_counts"]["train"])),
        ("validation_split", str(summary["split_counts"]["validation"])),
        ("test_split", str(summary["split_counts"]["test"])),
        ("precision", str(primary["precision"])),
        ("recall", str(primary["recall"])),
        ("f1", str(primary["f1"])),
        ("accuracy", str(primary["accuracy"])),
        ("roc_auc", str(primary["roc_auc"])),
        ("pr_auc", str(primary["pr_auc"])),
        ("tp", str(primary["tp"])),
        ("fp", str(primary["fp"])),
        ("fn", str(primary["fn"])),
        ("tn", str(primary["tn"])),
        ("baseline_f1", str(baseline["f1"])),
        ("baseline_recall", str(baseline["recall"])),
        ("ablation_f1", str(ablation["f1"])),
        ("ablation_recall", str(ablation["recall"])),
    ]


def main() -> int:
    metrics = load_json(RESULTS / "metrics.json")
    summary = load_json(ROOT / "data" / "processed" / "dataset_summary.json")
    report_path = ROOT / "Marianovskiy_VKR_ADAS_final_clean.docx"
    deck_path = ROOT / "Marianovskiy_VKR_ADAS_defense_final_clean.pptx"
    combined = ""
    if report_path.exists():
        combined += docx_text(report_path)
    if deck_path.exists():
        combined += "\n" + pptx_text(deck_path)

    rows = []
    missing = []
    for label, token in expected_tokens(metrics, summary):
        ok = token in combined
        rows.append((label, token, ok))
        if not ok:
            missing.append((label, token))

    lines = [
        "# RESULTS CONSISTENCY V2",
        "",
        f"Checked report: `{report_path.name}`",
        f"Checked presentation: `{deck_path.name}`",
        "",
        "| Check | Expected value | Found |",
        "|---|---:|---|",
    ]
    for label, token, ok in rows:
        lines.append(f"| {label} | {token} | {'yes' if ok else 'no'} |")
    lines.append("")
    lines.append("Status: passed" if not missing else "Status: failed")
    (ROOT / "RESULTS_CONSISTENCY_V2.md").write_text("\n".join(lines) + "\n", encoding="utf-8")

    if missing:
        print(json.dumps({"missing": missing}, ensure_ascii=False, indent=2))
        return 1
    print("Results consistency V2 passed")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
