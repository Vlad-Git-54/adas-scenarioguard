"""Build final VKR DOCX artifacts from repository results."""
from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Iterable, List, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
FIGURES = ROOT / "figures"
RESULTS = ROOT / "results"
FINAL = ROOT / "final"

AUTHOR = "Марьяновский Владислав Андреевич"
GROUP = "292405-1"
THEME = (
    "Разработка методов обнаружения и обработки редких и критических сценариев "
    "в мультимодальном восприятии систем ADAS для повышения безопасности "
    "в сложных погодных и дорожных условиях"
)


def metrics() -> dict:
    return json.loads((RESULTS / "metrics.json").read_text(encoding="utf-8"))


def dataset_summary() -> dict:
    return json.loads((ROOT / "data" / "processed" / "dataset_summary.json").read_text(encoding="utf-8"))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def apply_page_setup(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)


def clear_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    for paragraph in section.footer.paragraphs:
        paragraph.clear()


def set_page_number_start(section, start: int) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def enable_page_numbers(section, start: int | None = None) -> None:
    section.footer.is_linked_to_previous = False
    section.different_first_page_header_footer = False
    clear_footer(section)
    if start is not None:
        set_page_number_start(section, start)
    add_page_number(section.footer.paragraphs[0])


def setup_document(doc: Document, page_numbers: bool = True, first_page_number: int | None = None) -> None:
    section = doc.sections[0]
    apply_page_setup(section)
    if page_numbers:
        section.different_first_page_header_footer = True
        if first_page_number is not None:
            set_page_number_start(section, first_page_number)
        add_page_number(section.footer.paragraphs[0])
    else:
        section.different_first_page_header_footer = False
        clear_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)
    styles["Heading 1"].font.size = Pt(12)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.bold = True


def set_core_properties(doc: Document, title: str) -> None:
    props = doc.core_properties
    props.title = title
    props.subject = THEME
    props.author = AUTHOR
    props.comments = ""
    props.keywords = "ADAS, KITTI, critical scene, sensor fusion"
    props.last_modified_by = AUTHOR


def start_numbered_section(doc: Document, first_page_number: int) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_page_setup(section)
    enable_page_numbers(section, first_page_number)


def p(doc: Document, text: str = "", align=None, bold: bool = False, first_line: bool = True):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.first_line_indent = Cm(1.25) if first_line else Cm(0)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return para


def structural_heading(doc: Document, text: str, page_break: bool = True) -> None:
    if page_break:
        doc.add_page_break()
    para = p(doc, text.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_line=False)
    para.style = doc.styles["Heading 1"]


def h1(doc: Document, text: str) -> None:
    doc.add_page_break()
    para = doc.add_paragraph()
    para.style = doc.styles["Heading 1"]
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.add_run(text).bold = True


def h2(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.style = doc.styles["Heading 2"]
    para.paragraph_format.first_line_indent = Cm(1.25)
    para.add_run(text).bold = True


def table_caption(doc: Document, number: int, title: str) -> None:
    para = p(doc, f"Таблица {number} – {title}", first_line=False)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def figure_caption(doc: Document, number: int, title: str) -> None:
    para = p(doc, f"Рисунок {number} – {title}", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)


def add_table(doc: Document, rows: Sequence[Sequence[str]], widths: Sequence[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for i, text in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[i], text, bold=True)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], str(text))
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Cm(width)


def add_figure(doc: Document, filename: str, number: int, caption: str, width: float = 5.8) -> None:
    path = FIGURES / filename
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        figure_caption(doc, number, caption)


def add_title_page(doc: Document) -> None:
    for text in [
        "Министерство науки и высшего образования Российской Федерации",
        "НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ",
        "ТОМСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ (НИ ТГУ)",
        "Институт дистанционного образования",
        "Направление 09.04.03 Прикладная информатика",
        "Направленность «Компьютерное зрение и нейронные сети»",
    ]:
        p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    # Keep the title page on one A4 page with 14 pt body formatting.
    approval = [
        "ДОПУСТИТЬ К ЗАЩИТЕ В ГЭК",
        "Руководитель ОПОП",
        "Старший лаборант",
        "______________ А.В. Осинцев",
        "«___» ____________ 2026 г.",
    ]
    for text in approval:
        p(doc, text, align=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)
    p(doc, "", first_line=False)
    p(doc, "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА МАГИСТРА", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_line=False)
    p(doc, THEME, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_line=False)
    p(doc, "", first_line=False)
    p(doc, f"Автор работы: {AUTHOR}", first_line=False)
    p(doc, f"Группа: {GROUP}", first_line=False)
    p(doc, "Подпись автора: __________________________", first_line=False)
    p(doc, "Руководитель: ________________________________", first_line=False)
    p(doc, "Должность, ученая степень: ____________________", first_line=False)
    p(doc, "Подпись руководителя: __________________________", first_line=False)
    p(doc, "Дата: «___» ____________ 2026 г.", first_line=False)
    p(doc, "Томск 2026", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)

def add_annotation(doc: Document, m: dict, summary: dict, page_break: bool = True) -> None:
    structural_heading(doc, "АННОТАЦИЯ", page_break=page_break)
    primary = m["models"][m["primary_model"]]
    paragraphs = [
        f"В выпускной квалификационной работе рассматривается задача обнаружения редких и критических сценариев в мультимодальном восприятии ADAS. Работа сфокусирована не на полном автопилоте, а на слое надежности восприятия, который помогает понять, когда дорожная сцена требует осторожной обработки, повторной проверки или включения в набор сложных примеров.",
        f"В программной части реализован прототип ADAS ScenarioGuard. Основной эксперимент выполнен на реальных аннотациях KITTI Object Detection. Из {summary['num_scenes']} размеченной дорожной сцены построена scenario-level таблица признаков. Целевая переменная critical_scene получена фиксированным правилом из класса объекта, 3D-дистанции, бокового положения, occlusion и truncation, поскольку KITTI не содержит готовой метки критичности ADAS.",
        f"Для проверки обучены три модели logistic regression: baseline, proposed и ablation. Primary model proposed_reliability_logreg на test split из {primary['num_examples']} сцен получила precision = {primary['precision']}, recall = {primary['recall']}, F1 = {primary['f1']}, ROC AUC = {primary['roc_auc']} и PR AUC = {primary['pr_auc']}. Матрица ошибок содержит TP = {primary['tp']}, FP = {primary['fp']}, FN = {primary['fn']} и TN = {primary['tn']}.",
        "Практический результат работы представлен репозиторием с кодом подготовки данных, обучения, оценки, генерации графиков и диаграмм. Отдельно подготовлены разделы воспроизводимости, анализа ошибок, ограничения метода, финальная презентация защиты и предметный указатель компетенций.",
        "Ключевые слова: ADAS, мультимодальное восприятие, KITTI, critical scene, sensor fusion, uncertainty estimation, logistic regression, анализ ошибок, безопасность дорожной сцены.",
    ]
    for text in paragraphs:
        p(doc, text)


def add_toc_clean(doc: Document) -> None:
    structural_heading(doc, "ОГЛАВЛЕНИЕ")
    entries = [
        ('ВВЕДЕНИЕ', 7),
        ('1 Анализ предметной области и существующих подходов', 10),
        ('1.1 ADAS и мультимодальное восприятие', 10),
        ('1.2 Редкие критические сценарии и corner cases', 10),
        ('1.3 Плохая погода, окклюзия и деградация сенсоров', 11),
        ('1.4 Почему средние метрики не гарантируют безопасность', 11),
        ('1.5 OOD detection и uncertainty estimation', 12),
        ('1.6 Виды sensor fusion', 12),
        ('1.7 Датасеты и симуляторы', 13),
        ('1.8 Аналоги и ограничения известных решений', 13),
        ('1.9 Проверяемость исследований в ADAS', 14),
        ('1.10 Связь литературы с задачей ВКР', 14),
        ('1.11 Ограничения обзорной части', 16),
        ('1.12 Выводы по главе 1', 17),
        ('2 Методика обнаружения и обработки редких критических сценариев', 19),
        ('2.1 Классификация сценариев', 19),
        ('2.2 Признаки качества наблюдения', 19),
        ('2.3 Sensor reliability', 20),
        ('2.4 Uncertainty score', 20),
        ('2.5 Risk score', 20),
        ('2.6 Adaptive fusion', 21),
        ('2.7 Обработка отказа сенсора', 21),
        ('2.8 Метрики', 22),
        ('2.9 План эксперимента', 22),
        ('2.10 Этика и приватность', 23),
        ('2.11 Требования к промышленной проверке', 23),
        ('2.12 Формирование метки critical_scene', 24),
        ('2.13 Защита от утечки данных', 25),
        ('2.14 Интерпретация метрик', 27),
        ('2.15 Выводы по главе 2', 29),
        ('3 Реализация прототипа и экспериментальная проверка', 31),
        ('3.1 Структура репозитория', 31),
        ('3.2 Используемые инструменты', 31),
        ('3.3 Подготовка данных', 31),
        ('3.4 Обучение моделей', 32),
        ('3.5 Настройка порогов', 32),
        ('3.6 Результаты', 33),
        ('3.7 Анализ ошибок', 33),
        ('3.8 Ограничения', 33),
        ('3.9 Воспроизводимость', 34),
        ('3.10 Сравнение с литературным примером BEVFusion', 34),
        ('3.11 Проверка целостности результатов', 35),
        ('3.12 Перенос на raw sensor pipeline', 35),
        ('3.13 Подробный разбор FP и FN', 35),
        ('3.14 Репликация эксперимента', 37),
        ('3.15 Границы переноса на raw sensor pipeline', 39),
        ('ЗАКЛЮЧЕНИЕ', 47),
        ('ЛИТЕРАТУРА', 49),
        ('ПРИЛОЖЕНИЯ', 51),
    ]
    for title, page in entries:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.left_indent = Cm(0.7) if title[:2].count(".") else Cm(0)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.tab_stops.add_tab_stop(Cm(17), WD_TAB_ALIGNMENT.RIGHT)
        title_run = para.add_run(title)
        para.add_run("\t")
        page_run = para.add_run(str(page))
        for run in (title_run, page_run):
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_terms(doc: Document) -> None:
    structural_heading(doc, "ПЕРЕЧЕНЬ УСЛОВНЫХ ОБОЗНАЧЕНИЙ, СОКРАЩЕНИЙ И ТЕРМИНОВ")
    rows = [
        ["Термин", "Расшифровка"],
        ["ADAS", "Advanced Driver Assistance Systems, системы помощи водителю"],
        ["LiDAR", "сенсор, измеряющий расстояния по отражению лазерного излучения"],
        ["Radar", "радиолокационный сенсор для оценки расстояния и скорости"],
        ["Sensor fusion", "объединение данных нескольких сенсоров"],
        ["OOD", "out-of-distribution, ситуация вне привычного распределения данных"],
        ["FNR", "false negative rate, доля пропущенных критических сцен"],
        ["FPR", "false positive rate, доля безопасных сцен, ошибочно отмеченных как критические"],
        ["KITTI", "открытый датасет и benchmark для задач восприятия в автономном вождении"],
    ]
    add_table(doc, rows, widths=[4, 12])


def add_intro(doc: Document, m: dict, summary: dict) -> None:
    structural_heading(doc, "ВВЕДЕНИЕ")
    intro = [
        "Системы помощи водителю используют камеры, LiDAR, radar и алгоритмы анализа дорожной сцены для предупреждения об опасных ситуациях [3, 4]. На практике важно не только обнаружить объект, но и понять, можно ли доверять результату восприятия. В тумане, дожде, снегу, ночью, при бликах и частичной видимости объект может быть найден с низкой уверенностью или не найден вовсе [2]. Для ADAS такая ситуация опасна, потому что среднее качество модели не показывает поведение на редких критических сценах.",
        "Степень разработанности темы связана с несколькими направлениями. В работах по corner cases описываются редкие ситуации для восприятия в высокоавтоматизированном вождении [1]. В исследованиях adverse weather показано, что сенсоры деградируют по-разному [2]. Обзоры sensor fusion систематизируют раннее, промежуточное и позднее объединение модальностей [3, 4]. Отдельно развиваются OOD detection и uncertainty estimation, где модель должна оценить собственную ненадежность [11-13].",
        "Проблема работы состоит в том, что обычная модель восприятия может показать высокую среднюю точность, но пропускать редкие опасные сцены. В ADAS пропуск пешехода впереди автомобиля или близкого препятствия при плохой видимости имеет большую цену. Поэтому требуется отдельный слой оценки критичности, который работает поверх признаков сцены и выделяет случаи, требующие осторожного режима.",
        "Объект исследования: мультимодальное восприятие систем ADAS, включающее данные камеры, LiDAR, radar и алгоритмы анализа дорожной сцены. Предмет исследования: методы обнаружения и обработки редких и критических сценариев в мультимодальном восприятии ADAS, включая corner case detection, OOD detection, uncertainty estimation и adaptive sensor fusion.",
        "Цель работы: разработать и экспериментально проверить методику обнаружения и обработки редких критических сценариев в мультимодальном восприятии ADAS для повышения надежности восприятия в сложных погодных и дорожных условиях.",
        "Для достижения цели решены задачи: проанализирована предметная область ADAS, рассмотрены подходы sensor fusion и uncertainty estimation, сформирована классификация критических сценариев, разработаны признаки оценки надежности, реализован программный прототип, подготовлен реальный scenario-level набор на основе KITTI, обучены baseline и proposed модели, рассчитаны метрики, проведен анализ ошибок и сформулированы ограничения метода.",
        "Методы исследования включают анализ научных источников, инженерное проектирование признаков, обработку реальных аннотаций KITTI Object Detection, обучение logistic regression на NumPy, подбор порога на validation split, расчет precision, recall, F1, FNR, FPR, ROC AUC, PR AUC, анализ ошибок и подготовку воспроизводимых артефактов.",
        "Научная новизна состоит в разработке и апробации методики scenario-level оценки надежности мультимодального восприятия ADAS, в которой редкий критический сценарий рассматривается как сочетание опасного класса объекта, геометрической близости, частичной видимости, повышенной неопределенности и риска пропуска опасного объекта.",
        "Практическая значимость состоит в том, что прототип может использоваться как вспомогательный модуль тестирования ADAS. Он выделяет сцены, где результат восприятия требует осторожной обработки, повторной проверки или добавления в набор сложных примеров для дообучения. Прототип не является сертифицированной системой безопасности и не предназначен для самостоятельного управления автомобилем.",
        "Границы исследования заданы доступными данными и вычислительными ресурсами. Основной эксперимент выполнен на реальных аннотациях KITTI [5, 6], но без обучения нейросетевой модели по сырым изображениям и облакам точек. Текущая модель является табличной и воспроизводимо обучается на CPU; raw sensor pipeline относится к следующему, отдельному этапу исследования.",
        "Работа состоит из введения, трех глав, заключения, списка литературы и приложений. Первая глава описывает предметную область и известные подходы. Вторая глава формулирует методику оценки критических сценариев. Третья глава описывает реализацию, эксперимент, результаты, ошибки, ограничения и воспроизводимость.",
    ]
    for text in intro:
        p(doc, text)


def add_conclusion(doc: Document, m: dict, summary: dict) -> None:
    structural_heading(doc, "ЗАКЛЮЧЕНИЕ")
    primary = m["models"][m["primary_model"]]
    conclusions = [
        "В работе разработана и проверена методика обнаружения редких и критических сценариев в мультимодальном восприятии ADAS. Методика рассматривает критическую сцену как сочетание класса объекта, геометрической близости, положения относительно траектории, частичной видимости, proxy-качества сенсорных признаков и неопределенности.",
        "Первая задача выполнена через анализ предметной области ADAS и проблемы редких сценариев. Показано, что средние метрики детекции не дают достаточного ответа о безопасности в сложных сценах. Вторая задача выполнена через обзор sensor fusion, OOD detection, uncertainty estimation, датасетов KITTI, nuScenes и симулятора CARLA.",
        "Третья задача выполнена через классификацию сценариев по типу объекта, геометрии, окклюзии, truncation и риску пропуска. Четвертая задача выполнена через разработку признаков camera_quality_proxy, lidar_geometry_quality_proxy, uncertainty_proxy и risk_prior. Пятая задача выполнена через программный прототип ADAS ScenarioGuard.",
        f"Шестая задача выполнена через подготовку реального control dataset на основе KITTI Object Detection. Таблица содержит {summary['num_scenes']} сцен, train split содержит {summary['split_counts']['train']} сцен, validation split содержит {summary['split_counts']['validation']} сцен, test split содержит {summary['split_counts']['test']} сцен.",
        f"Седьмая и восьмая задачи выполнены через обучение и сравнение baseline, proposed и ablation моделей. Primary model получила precision = {primary['precision']}, recall = {primary['recall']}, F1 = {primary['f1']}, accuracy = {primary['accuracy']}, ROC AUC = {primary['roc_auc']} и PR AUC = {primary['pr_auc']}. FNR составил {primary['false_negative_rate']}, FPR составил {primary['false_positive_rate']}.",
        f"Девятая задача выполнена через анализ ошибок. На test split найдено FP = {primary['fp']} и FN = {primary['fn']}. False positive связаны с близкими объектами, окклюзией и высоким risk prior. False negative возникают там, где derived target относит сцену к критической, но совокупный score модели остается ниже порога.",
        "Десятая задача выполнена через описание ограничений. KITTI не содержит radar и готовую метку critical_scene, поэтому основной эксперимент не является промышленной валидацией мультимодальной ADAS. Модель работает на scenario-level признаках, а не на сырых изображениях и облаках точек. Для внедрения нужны реальные сенсорные данные, проверка отказов, сертификация и юридическая оценка.",
        "Разработанный прототип может применяться как исследовательский модуль тестирования ADAS. Он помогает выделять сцены, где восприятие требует осторожной обработки, повторной проверки или расширения обучающего набора. Дальнейшее развитие включает raw image и LiDAR pipeline, подключение radar, проверку на nuScenes и CARLA, а также отдельный выбор вычислительного backend для более тяжелых моделей.",
    ]
    for text in conclusions:
        p(doc, text)


def add_literature(doc: Document) -> None:
    structural_heading(doc, "ЛИТЕРАТУРА")
    sources = [
        "Heidecker F. An Application-Driven Conceptualization of Corner Cases for Perception in Highly Automated Driving / F. Heidecker, J. Breitenstein, K. Rösch [et al.] // 2021 IEEE Intelligent Vehicles Symposium. – 2021. – P. 644-651. – DOI: 10.1109/IV48863.2021.9575933.",
        "Bijelic M. Seeing Through Fog Without Seeing Fog: Deep Multimodal Sensor Fusion in Unseen Adverse Weather / M. Bijelic, T. Gruber, F. Mannan [et al.] // Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition. – 2020. – P. 11679-11689. – DOI: 10.1109/CVPR42600.2020.01170.",
        "Huang K. Multi-modal Sensor Fusion for Auto Driving Perception: A Survey / K. Huang, B. Shi, X. Li [et al.] // arXiv. – 2024. – arXiv:2202.02703.",
        "Feng D. Deep Multi-modal Object Detection and Semantic Segmentation for Autonomous Driving: Datasets, Methods, and Challenges / D. Feng, C. Haase-Schütz, L. Rosenbaum [et al.] // IEEE Transactions on Intelligent Transportation Systems. – 2021. – Vol. 22, № 3. – P. 1341-1360.",
        "Geiger A. Are we ready for Autonomous Driving? The KITTI Vision Benchmark Suite / A. Geiger, P. Lenz, R. Urtasun // IEEE Conference on Computer Vision and Pattern Recognition. – 2012. – P. 3354-3361.",
        "KITTI Vision Benchmark Suite [Электронный ресурс] // Karlsruhe Institute of Technology and Toyota Technological Institute at Chicago. – URL: https://www.cvlibs.net/datasets/kitti/ (дата обращения: 07.07.2026).",
        "Caesar H. nuScenes: A multimodal dataset for autonomous driving / H. Caesar, V. Bankiti, A. H. Lang [et al.] // IEEE/CVF Conference on Computer Vision and Pattern Recognition. – 2020. – P. 11621-11631.",
        "Dosovitskiy A. CARLA: An Open Urban Driving Simulator / A. Dosovitskiy, G. Ros, F. Codevilla [et al.] // Proceedings of the 1st Annual Conference on Robot Learning. – 2017. – P. 1-16.",
        "Liu Z. BEVFusion: Multi-Task Multi-Sensor Fusion with Unified Bird's-Eye View Representation / Z. Liu, H. Tang, A. Amini [et al.] // arXiv. – 2022. – arXiv:2205.13542.",
        "Kumar S. Evaluating the Impact of Weather-Induced Sensor Occlusion on BEVFusion for 3D Object Detection / S. Kumar, T. Brophy, E. M. Grua [et al.] // arXiv. – 2025. – URL: https://arxiv.org/abs/2511.04347 (дата обращения: 07.07.2026).",
        "Hendrycks D. A Baseline for Detecting Misclassified and Out-of-Distribution Examples in Neural Networks / D. Hendrycks, K. Gimpel // International Conference on Learning Representations. – 2017.",
        "Gal Y. Dropout as a Bayesian Approximation: Representing Model Uncertainty in Deep Learning / Y. Gal, Z. Ghahramani // International Conference on Machine Learning. – 2016. – P. 1050-1059.",
        "Lakshminarayanan B. Simple and Scalable Predictive Uncertainty Estimation using Deep Ensembles / B. Lakshminarayanan, A. Pritzel, C. Blundell // Advances in Neural Information Processing Systems. – 2017.",
        "ADAS ScenarioGuard: репозиторий проекта [Электронный ресурс]. – URL: https://github.com/Vlad-Git-54/adas-scenarioguard (дата обращения: 07.07.2026).",
    ]
    for i, source in enumerate(sources, 1):
        p(doc, f"{i}. {source}", first_line=False)


def add_appendices(doc: Document, m: dict) -> None:
    structural_heading(doc, "ПРИЛОЖЕНИЯ")
    appendices = [
        ("ПРИЛОЖЕНИЕ А", "Задание на ВКР", ["В приложении А приведено задание на выполнение выпускной квалификационной работы. Поля подписей и даты оставлены пустыми для заполнения по официальной процедуре."]),
        ("ПРИЛОЖЕНИЕ Б", "Архитектура прототипа", ["Архитектура состоит из подготовки KITTI label_2, построения scenario table, обучения logistic regression, оценки test split, генерации графиков и документирования результатов. Основные схемы приведены в главе 3 и сохранены в папке figures."]),
        ("ПРИЛОЖЕНИЕ В", "Примеры входных JSON-сцен", ["JSON demo сохраняет идею мультимодального confidence от camera, LiDAR и radar. Этот demo не является основным источником метрик, но помогает проверить CLI и объяснить работу risk score на понятной сцене."]),
        ("ПРИЛОЖЕНИЕ Г", "Ссылка на репозиторий и фрагменты кода", ["Репозиторий проекта: https://github.com/Vlad-Git-54/adas-scenarioguard. Основные файлы: src/adas_scenarioguard/experiment.py, scripts/prepare_data.py, scripts/train.py, scripts/evaluate.py."]),
        ("ПРИЛОЖЕНИЕ Д", "Таблицы метрик", [f"Primary model: {m['primary_model']}. Результаты сохранены в results/metrics.json и results/metrics.csv. Эти файлы формируются командой python scripts/evaluate.py."]),
        ("ПРИЛОЖЕНИЕ Е", "Анализ ошибок", ["Подробный список FP и FN находится в results/error_cases.csv и docs/error_analysis.md. Ошибки связаны с близкими объектами, окклюзией, truncation и недостаточным score для части положительных сцен."]),
        ("ПРИЛОЖЕНИЕ Ж", "Инструкция по воспроизведению", ["Полный цикл: python scripts/prepare_data.py, python scripts/train.py, python scripts/evaluate.py, python scripts/make_figures.py, python scripts/make_diagrams.py, python scripts/export_report_assets.py, python scripts/export_results.py, python -m pytest -q."]),
        ("ПРИЛОЖЕНИЕ И", "Предметный указатель компетенций", ["В приложении И приведен предметный указатель компетенций, соотнесенный со структурными элементами работы. Он является последним приложением к работе."]),
    ]
    for head, title, paragraphs in appendices:
        doc.add_page_break()
        p(doc, head, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_line=False)
        p(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_line=False)
        for text in paragraphs:
            p(doc, text)
    table_caption(doc, 6, "Предметный указатель компетенций")
    add_table(doc, competency_rows(), widths=[3, 8, 5])
    p(doc, "Руководитель ВКР: ______________________", first_line=False)


def competency_rows() -> List[List[str]]:
    return [
        ["Компетенция", "Проявление в работе", "Разделы ВКР"],
        ["УК-1", "постановка проблемы и выбор способа исследования", "Введение, 1.1-1.4"],
        ["УК-2", "планирование этапов проекта и проверяемых результатов", "Введение, 2.9, 3.9"],
        ["УК-3", "организация воспроизводимого набора файлов и ролей проверки", "3.1, 3.9, приложения"],
        ["УК-4", "подготовка текста, презентации и ответов на вопросы защиты", "Введение, заключение, презентация"],
        ["УК-5", "учет ограничений, этики и условий применения ADAS", "2.10, 3.8"],
        ["УК-6", "самостоятельная сборка и проверка итогового пакета", "3.9, приложения"],
        ["ОПК-1", "анализ предметной области и выбор научных источников", "Глава 1"],
        ["ОПК-2", "построение модели данных и формализация признаков", "2.1-2.5"],
        ["ОПК-3", "разработка программного прототипа и pipeline", "3.1-3.4"],
        ["ОПК-4", "расчет метрик и интерпретация результатов", "2.8, 3.6"],
        ["ОПК-5", "использование программных средств анализа данных", "3.1-3.5"],
        ["ОПК-6", "оценка надежности и ограничений модели", "3.7-3.8"],
        ["ОПК-7", "подготовка воспроизводимых материалов и документации", "3.9, приложения"],
        ["ОПК-8", "представление результатов в отчете и презентации", "Заключение, презентация"],
        ["ПК-1", "управление получением, хранением, передачей и обработкой больших данных на примере KITTI-таблицы", "3.3-3.7"],
        ["ПК-2", "оценка инфраструктуры анализа данных и направления ее развития", "3.8-3.12"],
    ]


def write_clean_sections(doc: Document, sections: Sequence[tuple[str, Sequence[str]]]) -> None:
    for heading, paragraphs in sections:
        h2(doc, heading)
        for text in paragraphs:
            p(doc, text)


def extend_section(sections: list[tuple[str, Sequence[str]]], marker: str, extra: Sequence[str]) -> None:
    for idx, (heading, paragraphs) in enumerate(sections):
        if marker in heading:
            sections[idx] = (heading, [*paragraphs, *extra])
            return
    raise ValueError(f"Section marker not found: {marker}")


def add_assignment_in_report_clean(doc: Document, page_break: bool = True) -> None:
    if page_break:
        doc.add_page_break()
        p(doc, "ЗАДАНИЕ ПО ВЫПОЛНЕНИЮ ВКР", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_line=False)
    else:
        p(doc, "ЗАДАНИЕ ПО ВЫПОЛНЕНИЮ ВКР", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_line=False)
    for text in [
        "Министерство науки и высшего образования Российской Федерации",
        "НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ",
        "ТОМСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ (НИ ТГУ)",
        "Институт дистанционного образования",
    ]:
        p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    p(doc, "", first_line=False)
    p(doc, "УТВЕРЖДАЮ", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, first_line=False)
    p(doc, "Руководитель ОПОП", align=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)
    p(doc, "Старший лаборант", align=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)
    p(doc, "______________ А.В. Осинцев", align=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)
    p(doc, "«___» ____________ 2026 г.", align=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)
    p(doc, "", first_line=False)
    p(doc, f"ЗАДАНИЕ по выполнению выпускной квалификационной работы магистра обучающемуся {AUTHOR}", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_line=False)
    assignment_lines = [
        f"1. Тема работы: {THEME}.",
        "2. Срок сдачи работы в Центр педагогического дизайна и онлайн-обучения ИДО ТГУ: «___» ____________ 2026 г. Срок сдачи работы в ГЭК: «___» ____________ 2026 г.",
        "3. Исходные данные к работе: реальные аннотации KITTI Object Detection, научные публикации по ADAS, corner cases, adverse weather, sensor fusion, OOD detection и uncertainty estimation, программный репозиторий с кодом подготовки данных, обучения, оценки и построения графиков.",
        "4. Объект исследования: мультимодальное восприятие систем ADAS, включая данные камеры, LiDAR, radar и алгоритмы анализа дорожной сцены.",
        "5. Предмет исследования: методы обнаружения и обработки редких критических сценариев в мультимодальном восприятии ADAS.",
        "6. Цель работы: разработать и экспериментально проверить методику обнаружения редких критических сценариев на реальных данных KITTI без использования синтетических обучающих сцен.",
        "7. Задачи работы: проанализировать предметную область, сформировать признаки качества наблюдения и риска, построить scenario-level таблицу, обучить baseline, proposed и ablation модели, рассчитать метрики, разобрать ошибки, подготовить отчет и презентацию.",
        "8. Методы исследования: анализ научной литературы, инженерное проектирование признаков, обработка аннотаций KITTI, logistic regression, расчет метрик классификации, анализ FP/FN случаев и проверка воспроизводимости.",
        "9. Отрасль применения: системы помощи водителю, тестирование ADAS, анализ редких дорожных сцен и подготовка наборов сложных примеров для дальнейшего обучения.",
        "10. Краткое содержание работы: обзор редких сценариев и деградации сенсоров, методика оценки reliability, uncertainty и risk, реализация прототипа, эксперимент на KITTI, анализ ошибок, ограничения и направление переноса на raw sensor pipeline.",
    ]
    for line in assignment_lines:
        p(doc, line, first_line=False)
    p(doc, "Обучающийся: ______________________", first_line=False)
    p(doc, "Руководитель ВКР: ______________________", first_line=False)
    p(doc, "Дата выдачи задания: «___» ____________ 2026 г.", first_line=False)


def add_chapter_1_clean(doc: Document) -> None:
    h1(doc, "1 Анализ предметной области и существующих подходов")
    sections = [
        (
            "1.1 ADAS и мультимодальное восприятие",
            [
                "Системы помощи водителю получают картину дорожной сцены из нескольких источников. Камера дает форму, цвет, разметку и контекст, LiDAR помогает оценивать геометрию, radar полезен для расстояния и скорости. На практике эти источники не заменяют друг друга полностью, потому что каждый сенсор имеет свои слабые места.",
                "Мультимодальное восприятие важно рассматривать не как простое сложение сигналов, а как задачу доверия к наблюдению. Одна и та же сцена может выглядеть безопасной по среднему качеству детекции, но стать опасной из-за частичной видимости пешехода, малой дистанции или слабой геометрии. Поэтому для ВКР выбран слой оценки сцены, который работает поверх признаков наблюдения и отдельно фиксирует риск.",
                "Основной эксперимент построен на реальных аннотациях KITTI Object Detection. Данные не заменяются синтетическими примерами и не дополняются вручную созданными погодными сценами. Такое ограничение делает результат более честным, но требует аккуратно отделять измеряемые признаки KITTI от будущего raw-sensor контура.",
            ],
        ),
        (
            "1.2 Редкие критические сценарии и corner cases",
            [
                "Опасные ситуации занимают малую долю дорожных наблюдений, но именно они определяют практическую цену ошибки. Для ADAS пропуск пешехода или близкого объекта впереди автомобиля намного серьезнее, чем лишнее предупреждение в спокойной сцене. Поэтому средняя точность модели не закрывает вопрос безопасности.",
                "Corner case в данной работе понимается как сцена, где привычные признаки надежности становятся недостаточными. К таким случаям относятся близкие уязвимые участники движения, частичная окклюзия, усечение объекта границей кадра и неоднозначная геометрия. Эти ситуации нельзя оценивать только по классу объекта, потому что контекст сцены меняет риск.",
                "В эксперименте критичность задается явным правилом из реальных KITTI-аннотаций. Правило учитывает класс объекта, положение впереди автомобиля, 3D-дистанцию, occlusion и truncation. Такая метка не является исходной меткой KITTI, но она воспроизводима и не зависит от ручной подгонки после обучения.",
            ],
        ),
        (
            "1.3 Плохая погода, окклюзия и деградация сенсоров",
            [
                "Дождь, туман, снег, ночь и блики снижают качество восприятия разными способами. Камера чувствительна к контрасту и освещению, LiDAR может страдать от отражений и дальности, radar сохраняет устойчивость в ряде погодных условий, но дает менее богатое описание формы. Эти различия объясняют, почему система должна оценивать надежность наблюдения, а не только факт наличия объекта.",
                "KITTI Object Detection не содержит полноценной разметки погодных режимов и radar-канала. Поэтому в основной модели используются только те признаки, которые можно извлечь из реальной разметки: occlusion, truncation, 2D-размер bounding box, 3D-дистанция и боковое положение. Это честное ограничение эксперимента, а не недостаток, скрытый в тексте.",
                "Окклюзия и усечение объекта полезны как proxy-признаки деградации наблюдения. Они не описывают всю физику сенсоров, но позволяют проверить идею reliability score на доступных данных. Для промышленной проверки потребуется набор с синхронными raw sensor данными, погодой, radar и сценариями отказов.",
            ],
        ),
        (
            "1.4 Почему средние метрики не гарантируют безопасность",
            [
                "Accuracy, ROC AUC и общая точность помогают сравнивать модели, но они сглаживают поведение на редких опасных случаях. Если критические сцены составляют меньшую часть набора, высокая accuracy может сохраняться даже при нежелательном числе пропусков. В ADAS такой результат нельзя считать достаточным.",
                "Работа делает акцент на recall, FNR и анализе FN-случаев. Recall показывает, какая доля критических сцен найдена, FNR показывает долю пропусков, а разбор ошибок объясняет, где именно модель недооценивает риск. Это ближе к инженерной задаче, чем одна итоговая средняя метрика.",
                "Precision также остается важной, потому что избыток ложных тревог снижает доверие к системе. Тем не менее приоритет в работе отдан снижению пропусков критических сцен. Такой выбор отражен в подборе порога на validation split и в сравнении primary model с baseline и ablation.",
            ],
        ),
        (
            "1.5 OOD detection и uncertainty estimation",
            [
                "Сцена вне привычного распределения может содержать знакомые классы объектов, но в непривычном сочетании расстояния, положения и качества наблюдения. OOD detection помогает заметить такие случаи, а uncertainty estimation показывает, насколько устойчивым является вывод модели. Для ADAS это разные, но связанные задачи.",
                "В магистерской работе не обучается отдельная нейросетевая модель uncertainty на изображениях. Вместо этого рассчитывается proxy uncertainty из признаков качества камеры и 3D-геометрии. Такой вариант проще, но воспроизводим на KITTI-аннотациях и подходит для проверки методики на уровне сцены.",
                "Неопределенность используется не как декоративная оценка, а как входной признак primary model. Если камера частично теряет объект или геометрия сцены становится ненадежной, uncertainty_proxy повышается и влияет на вероятность критической сцены. Это связывает теоретический обзор с экспериментом.",
            ],
        ),
        (
            "1.6 Виды sensor fusion",
            [
                "Раннее объединение работает с сырыми сигналами или низкоуровневыми признаками. Промежуточное объединение соединяет представления, полученные отдельными ветвями модели. Позднее объединение комбинирует уже готовые решения или confidence отдельных детекторов. Выбор уровня зависит от данных, вычислительных ресурсов и доступности синхронных сенсоров.",
                "Для полного raw fusion нужны изображения, облака точек, radar, калибровка и вычислительный контур обучения. В текущем эксперименте используется более узкий scenario-level fusion: признаки класса, геометрии, окклюзии, усечения, надежности и риска объединяются в табличной модели. Это не подменяет промышленную multimodal perception систему, но проверяет важную часть логики.",
                "JSON demo в проекте показывает идею объединения confidence камеры, LiDAR и radar на простых сценах. Эти демонстрационные примеры не входят в обучение primary model и не влияют на метрики KITTI. Такое разделение нужно, чтобы не смешивать проверяемый эксперимент с иллюстрацией будущего расширения.",
            ],
        ),
        (
            "1.7 Датасеты и симуляторы",
            [
                "KITTI удобен для воспроизводимого эксперимента, потому что содержит открытые аннотации объектов, 2D/3D-геометрию, occlusion и truncation. Для ADAS-анализа его ограничение состоит в отсутствии готовой метки critical_scene, radar-канала и полной погодной разметки. Поэтому результаты по KITTI нужно трактовать как проверку методики на реальных аннотациях, а не как сертификацию мультимодальной системы.",
                "nuScenes полезен как следующий источник данных, потому что содержит камеры, LiDAR и radar. CARLA может применяться для стресс-тестов редких условий, но синтетические сцены требуют отдельной проверки переноса на реальные данные. В финальном эксперименте синтетика не используется для обучения и расчета метрик.",
                "Такой выбор данных соответствует исследовательской стадии проекта. Сначала проверяется воспроизводимая модель на реальном открытом наборе, затем метод можно переносить на более богатые датасеты и на raw sensor pipeline. Это снижает риск завышенных заявлений в тексте ВКР.",
            ],
        ),
        (
            "1.8 Аналоги и ограничения известных решений",
            [
                "Современные модели 3D object detection и BEV-представления показывают сильные результаты на стандартных benchmark-задачах. Однако высокая метрика на benchmark не означает, что система одинаково надежна при деградации сенсоров, редких объектах и нестандартных сочетаниях признаков. Поэтому в литературе отдельно обсуждаются corner cases, adverse weather и оценка uncertainty.",
                "BEVFusion и близкие подходы используются в работе как литературный контекст. Собственный эксперимент не сравнивает primary model с BEVFusion и не заявляет превосходство над ней. Рисунок 1 показывает опубликованный пример влияния деградации сенсоров, а не результат выполненного здесь обучения.",
                "Ограничение известных решений для данной ВКР состоит в трудной воспроизводимости полного промышленного контура. Не всегда доступны исходные данные, веса, вычислительные ресурсы и единые протоколы редких сценариев. Поэтому работа делает упор на прозрачную табличную методику, где каждый признак и каждая метрика проверяются из файлов проекта.",
            ],
        ),
        (
            "1.9 Проверяемость исследований в ADAS",
            [
                "Проверяемость особенно важна в задачах безопасности. Текст, код, данные и графики должны ссылаться на одни и те же числа, иначе результат сложно защищать как инженерную работу. В проекте метрики берутся из results/metrics.json, а таблицы и презентация пересобираются из этого источника.",
                "Разделение train, validation и test фиксируется заранее. Порог выбирается на validation split, а финальные числа выводятся на test split. Такой порядок снижает риск подгонки и делает проверку честнее.",
                "Репозиторий содержит скрипты подготовки данных, обучения, оценки, построения графиков и генерации документов. Команды полного цикла позволяют повторить расчет и сопоставить результаты с текстом ВКР. Для воспроизводимой проверки это важнее, чем декларация о точности без подтверждаемых файлов.",
            ],
        ),
        (
            "1.10 Выводы по главе 1",
            [
                "Редкие критические сцены требуют отдельной оценки, потому что средние метрики плохо показывают цену пропуска опасного случая. Надежность ADAS-восприятия зависит не только от класса объекта, но и от дистанции, положения, окклюзии, усечения и уверенности в наблюдении.",
                "Реальные аннотации KITTI позволяют построить проверяемую scenario-level модель, но не закрывают весь промышленный контур. В наборе нет radar и готовой метки critical_scene, поэтому целевая переменная выводится явным правилом. Это ограничение прямо учитывается в методике и в интерпретации результатов.",
                "Дальнейшая глава переводит обзор в инженерную схему: классификацию сценариев, признаки качества наблюдения, reliability, uncertainty, risk score, adaptive fusion, метрики и план эксперимента.",
            ],
        ),
    ]
    conclusion = sections[-1]
    sections = sections[:-1] + [
        (
            "1.10 Связь литературы с задачей ВКР",
            [
                "Обзор литературы в работе используется не как отдельный справочный блок, а как основание для выбора признаков и границ эксперимента. Публикации о corner cases помогают объяснить, почему редкие ситуации нельзя оценивать только средним качеством детекции [1]. Исследования adverse weather показывают, что ухудшение одного сенсора не всегда одинаково отражается на других каналах восприятия [2].",
                "Обзоры sensor fusion важны для постановки, потому что в них разные способы объединения модальностей рассматриваются через момент слияния и тип доступной информации [3, 4]. Для ВКР из этого следует практический вывод: если нет полного raw sensor набора, нельзя честно заявлять обучение полноценной multimodal fusion модели. Поэтому эксперимент ограничен scenario-level признаками, а перенос к raw sensor pipeline вынесен в дальнейшую работу.",
                "KITTI выбран не потому, что он идеально описывает плохую погоду или radar, а потому что это открытый и хорошо документированный benchmark с реальными аннотациями объектов [5, 6]. Такая база подходит для воспроизводимой проверки табличной постановки. При этом текст отдельно фиксирует, что critical_scene не является готовой меткой KITTI и строится как производная переменная.",
                "nuScenes и CARLA рассматриваются как возможные направления расширения, а не как источники текущих метрик [7, 8]. В nuScenes есть современная мультимодальная структура с radar, а CARLA удобна для управляемой генерации условий. В основной эксперимент они не включены, чтобы не смешивать разные источники данных и не создавать видимость более широкой проверки.",
                "Публикации по BEVFusion нужны для объяснения влияния деградации сенсоров на качество 3D object detection [9, 10]. Эти результаты не подменяют собственный эксперимент. В работе они используются как аргумент в пользу признаков reliability и uncertainty, а не как сравнение разработанного прототипа с BEVFusion.",
                "Работы по OOD detection и uncertainty estimation задают язык для описания ненадежных наблюдений [11-13]. В текущем прототипе uncertainty является инженерным proxy-признаком, а не байесовской неопределенностью нейросети. Такая формулировка сохраняет связь с литературой и не завышает уровень реализованной модели.",
            ],
        ),
        (
            "1.11 Ограничения обзорной части",
            [
                "Литературные источники дают широкий контекст, но они не снимают необходимость собственной проверки на сохраненном split. Если в статье описана деградация сенсора или архитектура fusion, это еще не означает, что тот же эффект будет измерен в таблице KITTI. Поэтому в работе разделяются три уровня: опубликованные результаты, выбранная методика и собственные метрики.",
                "Часть источников относится к нейросетевым системам восприятия, а прототип ВКР использует logistic regression. Это различие не скрывается. Нейросетевые статьи помогают описать проблему и будущий raw sensor этап, а табличная модель нужна для воспроизводимой проверки выбранных признаков в пределах доступных аннотаций.",
                "Понятия reliability, uncertainty и risk в тексте имеют инженерный смысл. Они не объявляются физическими параметрами сенсора и не заменяют измерение реальной камеры, LiDAR или radar. Такая осторожность важна, потому что KITTI label_2 хранит аннотации объектов, а не полный журнал сенсорных отказов.",
                "Обзор также ограничен задачей ВКР. В нем не рассматриваются вопросы управления автомобилем, планирования траектории, сертификации ASIL и юридической ответственности. Эти темы важны для промышленного ADAS, но они не входят в проверяемую часть работы.",
                "Список источников подобран так, чтобы закрыть основные опоры текста: редкие сценарии, сложная погода, sensor fusion, датасеты, симуляторы, BEVFusion, OOD и uncertainty. Каждая из этих групп связана с конкретным решением в методике или с ограничением текущего эксперимента.",
                "После такого обзора дальнейшая часть работы не должна повторять литературные формулировки. Задача методической главы состоит в том, чтобы перевести эти идеи в признаки, правила split, метрики и проверяемый pipeline.",
            ],
        ),
        ("1.12 Выводы по главе 1", conclusion[1]),
    ]
    write_clean_sections(doc, sections)
    table_caption(doc, 1, "Сравнение направлений, использованных в работе")
    add_table(doc, [
        ["Направление", "Что дает для ВКР", "Ограничение"],
        ["Corner cases", "помогает описать редкие опасные сцены", "нет единой универсальной метки"],
        ["Adverse weather", "объясняет деградацию разных сенсоров", "требует специальных данных"],
        ["Sensor fusion", "объединяет признаки разных источников", "ошибка калибровки может ухудшить результат"],
        ["OOD и uncertainty", "показывает ненадежность вывода", "оценка зависит от выбранной модели"],
    ], widths=[4, 7, 5])
    add_figure(doc, "bevfusion_literature_chart.png", 1, "Литературный пример влияния деградации сенсоров на BEVFusion по данным Kumar et al., 2025", 5.8)


def add_chapter_2_clean(doc: Document) -> None:
    h1(doc, "2 Методика обнаружения и обработки редких критических сценариев")
    sections = [
        (
            "2.1 Классификация сценариев",
            [
                "Единицей анализа выбрана дорожная сцена, а не отдельный bounding box. Такой выбор соответствует задаче ADAS: система должна понять, требует ли вся ситуация осторожного режима, повторной проверки или добавления в набор сложных примеров.",
                "Целевая переменная critical_scene вычисляется по явному правилу. Сцена считается критической, если уязвимый участник находится впереди автомобиля на дистанции до 35 м, транспортное средство впереди находится на дистанции до 12 м, объект с occlusion не ниже 2 находится впереди на дистанции до 24 м или объект с truncation не ниже 0,55 находится впереди на дистанции до 28 м.",
                "Такая классификация не выдается за исходный benchmark KITTI. Она нужна для воспроизводимого исследования редких сцен на доступных реальных аннотациях. Все пороги заданы в коде и не меняются после просмотра test split.",
            ],
        ),
        (
            "2.2 Признаки качества наблюдения",
            [
                "Базовые признаки описывают количество объектов, минимальную 3D-дистанцию, число объектов в переднем коридоре, число уязвимых участников, число близких объектов и максимальную нормированную площадь bounding box. Эти признаки дают грубую картину сложности сцены.",
                "Расширенные признаки добавляют occluded_count, truncated_count, mean_occlusion, mean_truncation, camera_quality_proxy, lidar_geometry_quality_proxy, uncertainty_proxy, risk_prior, max_abs_alpha и min_lateral_abs_m. Они позволяют модели учитывать не только наличие объекта, но и качество наблюдения.",
                "Все значения берутся из KITTI label_2 и производных расчетов. В таблицу признаков не добавляются синтетические погодные метки и ручные оценки после обучения. Это делает эксперимент уже, но честнее.",
            ],
        ),
        (
            "2.3 Sensor reliability",
            [
                "Reliability в работе трактуется как расчетная надежность наблюдения, доступная на уровне аннотаций. Для камеры используется proxy через truncation и occlusion, потому что усеченный или перекрытый объект хуже наблюдается визуально. Для 3D-геометрии используется proxy через дистанцию и occlusion.",
                "KITTI не содержит radar в используемом наборе, поэтому radar не входит в обучение primary model. Упоминание radar относится к общей постановке ADAS и к демонстрационному JSON-контуру, где показана идея объединения confidence. В метриках KITTI radar не участвует.",
                "Формально camera_quality_proxy снижается при росте максимального truncation и occlusion. lidar_geometry_quality_proxy снижается при ухудшении геометрических условий и окклюзии. Эти признаки не заменяют физическую модель сенсоров, но позволяют проверить, дает ли учет надежности выигрыш в recall и F1.",
            ],
        ),
        (
            "2.4 Uncertainty score",
            [
                "Uncertainty_proxy строится из деградации camera_quality_proxy и lidar_geometry_quality_proxy. Чем хуже наблюдение и геометрия, тем выше неопределенность. В коде используется фиксированная формула с весами 0,55 для camera degradation и 0,45 для geometry degradation.",
                "Значение uncertainty не является нейросетевой апостериорной неопределенностью. Это инженерный proxy-признак, который можно полностью воспроизвести из аннотаций. Его роль состоит в том, чтобы добавить в модель информацию о ненадежности сцены.",
                "Такой подход соответствует доступным данным. Он не требует синтетической генерации тумана, дождя или засветки и не создает видимость более богатого датасета, чем есть на самом деле.",
            ],
        ),
        (
            "2.5 Risk score",
            [
                "Risk_prior оценивает потенциальную опасность объекта до обучения модели. В расчет входят класс объекта, дистанция, положение по боковой оси, occlusion и truncation. Уязвимые участники движения получают больший базовый вклад, чем обычные транспортные средства.",
                "Distance risk растет при приближении объекта, lateral risk растет при попадании объекта в передний коридор, occlusion и truncation повышают риск из-за ухудшения наблюдения. Затем для сцены берется максимальный объектный риск. Это соответствует интуиции ADAS, где один опасный объект может сделать критической всю сцену.",
                "Risk_prior не является готовым решением. Он передается в proposed model вместе с остальными признаками, а итоговая вероятность критической сцены обучается на train split. Такой порядок отделяет экспертную структуру признаков от статистической настройки модели.",
            ],
        ),
        (
            "2.6 Adaptive fusion",
            [
                "Adaptive fusion в рамках ВКР реализован на уровне признаков. Модель получает базовые признаки сцены, признаки деградации наблюдения, uncertainty_proxy и risk_prior. Если наблюдение хуже, это влияет на итоговый score через обученные веса logistic regression.",
                "В отличие от raw sensor fusion, такой подход не объединяет изображения, облака точек и radar-сигналы напрямую. Он рассчитан на доступный набор KITTI и хорошо проверяется в репозитории. Для магистерской работы это разумная граница между исследовательской идеей и воспроизводимой реализацией.",
                "Демонстрационный JSON-контур дополняет методику примером weighted fusion confidence. Он нужен для иллюстрации будущего расширения, но не подмешивается в обучение KITTI-модели и не влияет на итоговые значения precision, recall, F1, ROC AUC и PR AUC.",
            ],
        ),
        (
            "2.7 Обработка отказа сенсора",
            [
                "Отказ сенсора в промышленной ADAS-системе должен переводить восприятие в осторожный режим. Примеры таких отказов: отсутствующий канал, резкое падение confidence, конфликт между сенсорами, сильная окклюзия или невозможность надежно оценить геометрию.",
                "В текущем эксперименте отказ не моделируется синтетическими изображениями. На KITTI-аннотациях используются признаки, которые отражают неполноту наблюдения: occlusion, truncation и ухудшение proxy-качества. Это ограничение не скрывается и отдельно учитывается при описании переноса на raw sensor pipeline.",
                "Будущее расширение с обучением моделей на изображениях и облаках точек потребует отдельного backend и отдельной проверки вычислительного контура. В финальных метриках текущей работы этот будущий контур не участвует, потому что обучение выполнено для табличной scenario-level модели.",
            ],
        ),
        (
            "2.8 Метрики",
            [
                "Precision показывает долю верных предупреждений среди всех предупреждений модели. Recall показывает долю найденных критических сцен. F1 объединяет precision и recall, а FNR отдельно показывает долю пропущенных критических сцен. В ADAS FNR особенно важен, потому что пропуск опасной сцены имеет высокую цену.",
                "ROC AUC и PR AUC используются как дополнительные интегральные показатели. Они помогают сравнивать модели при разных порогах, но не заменяют анализ выбранного порога и матрицы ошибок. Поэтому результаты представлены вместе с TP, FP, FN и TN.",
                "Порог primary model выбирается на validation split. Test split используется только для финального расчета. Такой порядок нужен, чтобы итоговые числа отражали обобщение модели, а не настройку под тест.",
            ],
        ),
        (
            "2.9 План эксперимента",
            [
                "Эксперимент начинается с загрузки и распаковки KITTI label_2. Затем каждая сцена переводится в строку scenario table, где фиксируются признаки объектов и целевая переменная critical_scene. После этого выполняется стратифицированное разделение на train, validation и test с seed 54.",
                "На train split обучаются три модели: baseline_kitti_logreg, proposed_reliability_logreg и ablation_without_3d_geometry. Baseline использует базовые признаки сцены. Proposed добавляет признаки качества наблюдения, uncertainty и risk. Ablation проверяет, что происходит при удалении части 3D-геометрии.",
                "После выбора порога на validation split выполняется оценка на test split. Скрипты сохраняют metrics.json, predictions, error cases, графики и промежуточные сведения о признаках. Рисунки 2 и 3 показывают pipeline обработки и основные сценарии использования прототипа.",
            ],
        ),
        (
            "2.10 Этика и приватность",
            [
                "В работе используется открытый исследовательский набор KITTI. Эксперимент не собирает новые персональные данные и не идентифицирует людей на изображениях. Все расчеты выполняются на уровне аннотаций объектов и производных признаков.",
                "Ограничение этической интерпретации состоит в том, что исследовательский прототип не должен использоваться как готовая система принятия решений на дороге. Он показывает способ выделения рискованных сцен для анализа, тестирования и последующей инженерной проверки.",
                "При переносе на промышленный контур потребуется проверка политики хранения данных, анонимизации, журналирования отказов и юридической ответственности. Эти вопросы выходят за рамки эксперимента KITTI, но должны быть включены в дальнейшую работу.",
            ],
        ),
        (
            "2.11 Требования к промышленной проверке",
            [
                "Промышленная проверка должна включать синхронные камеры, LiDAR, radar, погодные режимы, ночные сцены, реальные отказы сенсоров и контроль качества калибровки. Одного табличного эксперимента на KITTI недостаточно для вывода о готовности системы к эксплуатации.",
                "Нужны отдельные наборы для validation, test, стресс-тестов и regression checks. Метрики должны считаться не только по всему набору, но и по группам: близкие объекты, уязвимые участники, окклюзия, сильное усечение, дальние объекты и сложная геометрия.",
                "Для сертификационного уровня также требуются safety case, traceability требований, проверка отказов, журналирование версий данных и независимый аудит. ВКР не закрывает эти требования, но формирует воспроизводимую основу для следующего этапа.",
            ],
        ),
        (
            "2.12 Выводы по главе 2",
            [
                "Методика связывает предметную область с проверяемым экспериментом. Критическая сцена определяется явным правилом из реальных KITTI-аннотаций, признаки качества наблюдения и риска рассчитываются воспроизводимо, а модель обучается на фиксированном train split.",
                "Главная граница метода состоит в scenario-level постановке. Она не заменяет raw sensor fusion и не использует radar в KITTI-метриках. При этом она позволяет проверить, улучшают ли признаки reliability, uncertainty и risk обнаружение редких критических сцен.",
                "Следующая глава описывает программную реализацию, структуру репозитория, обучение моделей, результаты на test split и ограничения переноса на промышленный ADAS-контур.",
            ],
        ),
    ]
    conclusion = sections[-1]
    sections = sections[:-1] + [
        (
            "2.12 Формирование метки critical_scene",
            [
                "KITTI Object Detection не содержит исходного признака, который прямо называл бы сцену критической для ADAS. Поэтому целевая переменная строится как derived target. Это не ручная экспертная разметка после просмотра результатов, а фиксированное правило, примененное одинаково ко всем сценам до обучения модели [5, 6].",
                "Правило использует только признаки, доступные в label_2: класс объекта, 3D-дистанцию, боковое положение, occlusion и truncation. Уязвимый участник впереди автомобиля относится к критическим сценам при более длинной дистанции, потому что цена пропуска пешехода или велосипедиста выше. Близкий транспорт впереди получает отдельное условие из-за риска столкновения.",
                "Окклюзия и truncation включены в правило не как декоративные признаки, а как сигналы неполной видимости. Если объект частично закрыт или выходит за границу кадра, модель восприятия может получить меньше информации о его форме и положении. В ADAS такая ситуация должна попадать в разбор даже тогда, когда класс объекта распознан.",
                "Пороговые значения выбраны до оценки test split и сохранены в коде подготовки данных. Это защищает эксперимент от подгонки под итоговые метрики. При повторном запуске тот же набор аннотаций дает ту же целевую переменную, если не менять правило и seed разбиения.",
                "Derived target не равен промышленной оценке опасности. Он описывает учебную и исследовательскую постановку, пригодную для проверки признаков reliability, uncertainty и risk. В тексте поэтому используется формулировка critical_scene label, полученная из реальных KITTI-аннотаций, а не исходная метка benchmark.",
                "Такой подход делает эксперимент проверяемым. Любой спор о критичности сцены можно свести к конкретному условию правила и конкретной строке scenario table. Это лучше, чем использовать неявную ручную оценку, которую нельзя повторить.",
            ],
        ),
        (
            "2.13 Защита от утечки данных",
            [
                "Утечка данных в этой работе могла бы появиться в трех местах: при построении признаков, при стандартизации и при выборе порога. Поэтому pipeline устроен так, чтобы test split не участвовал в обучении и настройке. Таблица сцен строится целиком, но статистические параметры модели рассчитываются только по train split.",
                "Стратифицированное разбиение сохраняется в отдельном JSON-файле. Это позволяет повторять эксперимент без нового случайного распределения сцен. Seed 54 фиксирует состав train, validation и test, а не только итоговые числа в отчете.",
                "Стандартизация признаков обучается на train split и затем применяется к validation и test. Если бы средние значения и дисперсии считались по всему набору, test split частично влиял бы на обучение. Для маленьких и несбалансированных групп такая ошибка может заметно исказить оценку.",
                "Порог классификации выбирается на validation split. Test split используется только после выбора порога, когда модель и правило принятия решения уже зафиксированы. Это особенно важно для recall и FNR, потому что небольшая подстройка порога может уменьшить число FN именно на тестовых сценах.",
                "Файлы results/models.json, results/threshold_selection.md и results/metrics.json оставляют проверяемый след. По ним можно увидеть, какие признаки использованы, какой порог выбран и какие метрики получены на test split. Отчет и презентация берут числа из этих файлов.",
                "Защита от утечки данных не делает модель промышленной, но делает учебный эксперимент честным. Проверяющий может повторить расчет и получить те же размеры split, те же веса модели и ту же матрицу ошибок.",
            ],
        ),
        (
            "2.14 Интерпретация метрик",
            [
                "Precision и recall в задаче критических сцен читаются несимметрично. Ложное предупреждение ухудшает удобство и может перегружать downstream-модуль, но пропуск критической сцены опаснее. Поэтому в работе отдельно обсуждаются FN и FNR, а не только F1.",
                "F1 полезен как компактная сводная метрика, но он скрывает цену ошибки. Две модели могут иметь близкий F1 и разный recall. Для ADAS предпочтительнее модель, которая при разумном precision уменьшает число пропущенных критических сцен.",
                "Accuracy в таком наборе не должна быть главным показателем. Отрицательных сцен больше, чем положительных, поэтому высокая accuracy может сохраняться даже при заметном числе FN. В отчете accuracy оставлена для полноты, а выводы строятся через recall, F1, PR AUC и матрицу ошибок.",
                "ROC AUC показывает способность модели ранжировать сцены при разных порогах. PR AUC лучше отражает качество на положительном классе, когда критические сцены занимают меньшую долю. Поэтому обе кривые приведены вместе с выбранным порогом.",
                "Матрица ошибок переводит метрики в понятные случаи. TP показывает найденные критические сцены, FP показывает лишние предупреждения, FN показывает пропуски, TN показывает корректно спокойные сцены. Для обсуждения безопасности FN важнее среднего score.",
                "Такая интерпретация нужна для инженерной оценки результатов. Работа не обещает готовую систему предотвращения аварий, но показывает, как выбранные признаки меняют recall, F1 и количество пропущенных сцен на фиксированном test split.",
            ],
        ),
        ("2.15 Выводы по главе 2", conclusion[1]),
    ]
    extend_section(sections, "2.12", [
        "В практической реализации правило critical_scene хранится не как скрытая экспертная договоренность, а как часть кода подготовки данных. Это важно для проверки: если проверяющий или другой исследователь заново выполнит подготовку KITTI-таблицы, он увидит не только итоговую метку, но и условия, из которых она получена.",
        "Пороговые значения выбраны как инженерная аппроксимация близости и ухудшенной видимости, а не как результат подбора по test split. В работе поэтому не утверждается, что именно эти пороги являются универсальными для всех ADAS-систем. Они задают проверяемую постановку, на которой можно сравнить baseline и proposed модели.",
        "Отдельно учитываются уязвимые участники движения, потому что пропуск пешехода или велосипедиста имеет иную цену, чем ошибка на дальнем транспортном средстве. Такая логика делает derived target ближе к задаче безопасности, чем простая классификация по количеству объектов.",
        "Occlusion и truncation включены в правило как признаки неполного наблюдения. Если объект виден частично, система восприятия может ошибиться не только в классе, но и в геометрии; поэтому такие сцены полезно рассматривать как кандидаты на критичность даже при умеренной дистанции.",
        "При этом derived target не претендует на юридическую оценку дорожной опасности. Он используется как учебная и исследовательская метка, позволяющая честно сравнить модели на одном и том же наборе реальных сцен.",
        "Такое описание защищает работу от завышенного вывода: модель обучена не на синтетике и не на экспертной базе аварийных ситуаций, а на воспроизводимой интерпретации открытых аннотаций KITTI.",
    ])
    extend_section(sections, "2.13", [
        "Дополнительный источник утечки связан с повторной генерацией графиков и документов. Если в тексте вручную исправлять числа после просмотра PDF, можно случайно оставить старое значение в таблице или на слайде. Поэтому итоговые материалы собираются из `results/metrics.json` и CSV-файлов, а не из независимых ручных копий.",
        "Validation split используется только для выбора рабочего порога. Даже если на test split видно, что другой порог дал бы меньше FN, он не должен подменять уже выбранное решение. Иначе оценка перестанет быть финальной проверкой и станет продолжением настройки.",
        "Файл `data/processed/kitti_split.json` фиксирует состав train, validation и test. Это особенно важно для небольших групп редких сцен: случайное переразбиение может дать похожую среднюю accuracy, но заметно изменить FN и FP.",
        "Стандартизация признаков также хранится вместе с моделью. Если пересчитать средние и дисперсии на всех данных, модель получит информацию о test split косвенно, даже без доступа к целевым меткам.",
        "В отчете отдельно описан порядок error analysis. Ошибки рассматриваются после расчета метрик и не возвращаются в обучение текущей версии. Такой порядок ближе к реальной инженерной проверке, где тестовый результат сначала фиксируют, а затем разбирают причины.",
        "Сохранение predictions и error cases делает результат проверяемым. Можно открыть конкретную сцену, посмотреть score и признаки, а не ограничиваться одной итоговой строкой метрик.",
    ])
    extend_section(sections, "2.14", [
        "Для задачи критических сцен полезно смотреть не только точку выбранного порога, но и форму ROC и Precision-Recall кривых. Если кривая резко падает около рабочего диапазона, система чувствительна к небольшому изменению порога и требует дополнительной проверки.",
        "PR AUC важна потому, что положительный класс меньше отрицательного. В такой ситуации ROC AUC может выглядеть высокой, даже когда precision на рабочем пороге оставляет много ложных предупреждений.",
        "Матрица ошибок нужна для интерпретации результатов и последующей разработки. Она показывает, что числа в тексте не являются абстрактными метриками, и помогает сразу увидеть, где цена ошибки выше.",
        "FN в отчете рассматриваются отдельно от общей accuracy. Если модель пропускает критические сцены, высокая accuracy не компенсирует этот риск, потому что безопасных сцен в наборе статистически больше.",
        "FP также нельзя просто назвать безвредными. Слишком частые ложные предупреждения могут перегрузить следующий модуль или оператора, поэтому precision оставлен в качестве обязательной метрики рядом с recall.",
        "Итоговая интерпретация строится на балансе: proposed model должна повышать recall и F1 без неконтролируемого роста FP, а все выводы должны быть привязаны к фиксированному KITTI test split.",
    ])
    write_clean_sections(doc, sections)
    table_caption(doc, 2, "Формулы метрик и расчетных признаков")
    add_table(doc, [
        ["Показатель", "Формула", "Назначение"],
        ["Precision", "TP / (TP + FP)", "доля верных предупреждений среди всех предупреждений"],
        ["Recall", "TP / (TP + FN)", "доля найденных критических сцен"],
        ["F1", "2 · Precision · Recall / (Precision + Recall)", "сводная метрика баланса precision и recall"],
        ["FNR", "FN / (FN + TP)", "доля пропущенных критических сцен"],
        ["FPR", "FP / (FP + TN)", "доля безопасных сцен, отмеченных как критические"],
        ["Sensor reliability", "1 - degradation_proxy", "оценка качества наблюдения источника"],
        ["Uncertainty score", "0,55 · camera_degradation + 0,45 · geometry_degradation", "оценка ненадежности признаков сцены"],
        ["Risk score", "class + distance + lateral + occlusion + truncation", "априорная оценка опасности сцены"],
    ], widths=[3.5, 6.5, 6])
    add_figure(doc, "pipeline_diagram.png", 2, "Pipeline обработки сцены в ADAS ScenarioGuard", 6.0)
    add_figure(doc, "use_case_diagram.png", 3, "Сценарии использования прототипа", 6.0)


def add_chapter_3_clean(doc: Document, m: dict, summary: dict) -> None:
    h1(doc, "3 Реализация прототипа и экспериментальная проверка")
    primary = m["models"][m["primary_model"]]
    baseline = m["models"]["baseline_kitti_logreg"]
    ablation = m["models"]["ablation_without_3d_geometry"]
    sections = [
        (
            "3.1 Структура репозитория",
            [
                "Репозиторий разделен на код, данные, результаты, графики, документы и тесты. Такое разделение нужно для повторной сборки эксперимента: подготовка данных выполняется отдельно от обучения, оценка отдельно от генерации отчетных материалов.",
                "Основная логика находится в src/adas_scenarioguard. Скрипты в папке scripts запускают подготовку KITTI-таблицы, обучение, оценку, построение графиков, диаграмм и документов. Результаты сохраняются в results, а визуальные материалы в figures.",
                "Таблица 3 фиксирует назначение ключевых каталогов. Она нужна не как формальная опись файлов, а как карта воспроизводимости: по ней можно проверить, откуда берутся числа и изображения в тексте ВКР.",
            ],
        ),
        (
            "3.2 Используемые инструменты",
            [
                "Основной стек намеренно оставлен компактным. Python и NumPy используются для подготовки данных, обучения logistic regression и расчета метрик. Matplotlib строит графики, python-docx и инструменты презентаций формируют финальные материалы.",
                "Текущий эксперимент выполняется в CPU-friendly контуре, потому что модель обучается на табличных признаках scenario-level. Будущая версия с raw image, LiDAR и radar pipeline потребует отдельного вычислительного контура и не должна смешиваться с уже полученными метриками.",
                "Минимальный стек повышает воспроизводимость. Проверяющий может повторить расчет на Windows без сложной инфраструктуры и сравнить результаты с файлами results/metrics.json, results/models.json и figures.",
            ],
        ),
        (
            "3.3 Подготовка данных",
            [
                f"Скрипт подготовки данных формирует scenario table из {summary['num_scenes']} сцен KITTI Object Detection. Положительных critical_scene получилось {summary['positive_critical_scenes']}, отрицательных сцен {summary['negative_scenes']}. Разделение содержит {summary['split_counts']['train']} сцен в train, {summary['split_counts']['validation']} в validation и {summary['split_counts']['test']} в test.",
                "Каждая строка таблицы описывает одну сцену. Из аннотаций берутся классы объектов, 2D/3D-геометрия, occlusion, truncation, alpha и положение относительно автомобиля. Затем рассчитываются признаки качества наблюдения, uncertainty_proxy и risk_prior.",
                "Синтетические данные в обучение не добавлялись. JSON-сцены из папки data/samples используются для демонстрации CLI и не входят в расчет метрик KITTI.",
            ],
        ),
        (
            "3.4 Обучение моделей",
            [
                "Обучаются три logistic regression модели. Baseline использует базовые признаки дорожной сцены. Proposed добавляет reliability, uncertainty, risk и 3D-геометрию. Ablation исключает часть геометрических признаков, чтобы проверить их вклад.",
                "Перед обучением признаки стандартизируются по train split. Параметры стандартизации, веса моделей и выбранные пороги сохраняются в results/models.json. Такой порядок нужен, чтобы test split не участвовал в настройке.",
                "Primary model выбрана как proposed_reliability_logreg, потому что она дает лучший баланс recall и F1 среди проверенных вариантов. Это видно в таблице 4 и на рисунке 7.",
            ],
        ),
        (
            "3.5 Настройка порогов",
            [
                f"Порог primary model равен {primary['threshold']}. Он выбран на validation split, где сравнивались значения F1 и recall. Test split оставался закрытым для финального расчета.",
                "Такой порядок особенно важен для редких критических сцен. Если подбирать порог на test split, итоговый F1 может выглядеть лучше, но перестанет быть честной оценкой обобщения.",
                "В results/threshold_selection.md сохранена логика выбора порогов. Это позволяет увидеть, что итоговые числа в тексте не появились после ручной подгонки.",
            ],
        ),
        (
            "3.6 Результаты",
            [
                f"На test split из {primary['num_examples']} сцен primary model получила precision = {primary['precision']}, recall = {primary['recall']}, F1 = {primary['f1']}, accuracy = {primary['accuracy']}, ROC AUC = {primary['roc_auc']} и PR AUC = {primary['pr_auc']}. Эти значения совпадают с results/metrics.json.",
                f"Baseline показала F1 = {baseline['f1']} и recall = {baseline['recall']}. Ablation без части 3D-геометрии показала F1 = {ablation['f1']} и recall = {ablation['recall']}. Сравнение подтверждает, что признаки надежности и геометрии улучшают обнаружение критических сцен.",
                f"Матрица ошибок primary model содержит TP = {primary['tp']}, FP = {primary['fp']}, FN = {primary['fn']} и TN = {primary['tn']}. Для ADAS особенно важны FN = {primary['fn']}, потому что это пропущенные критические сцены. Рисунки 6-9 показывают матрицу ошибок, сравнение моделей, ROC и Precision-Recall кривые.",
            ],
        ),
        (
            "3.7 Анализ ошибок",
            [
                "False positive чаще возникают в сценах, которые похожи на опасные из-за близкого объекта, высокого risk_prior или частичной видимости. Такие ошибки неприятны, но обычно безопаснее, чем пропуск критической сцены.",
                "False negative появляются там, где derived target относит сцену к критической, но совокупный score модели остается ниже порога. Для таких случаев важны дистанция, боковое положение, окклюзия и вклад uncertainty_proxy.",
                "Список FP и FN сохранен в results/error_cases.csv, а разбор ошибок вынесен в docs/error_analysis.md. Рисунок 10 показывает поведение по группам сцен, рисунок 11 показывает вклад признаков через ablation.",
            ],
        ),
        (
            "3.8 Ограничения",
            [
                "Главное ограничение связано с данными. KITTI Object Detection не содержит radar и не имеет исходной метки ADAS critical_scene. Поэтому целевая переменная является derived target, полученным из реальных аннотаций по фиксированному правилу.",
                "Модель работает на scenario-level признаках, а не на сырых изображениях и облаках точек. Она не заменяет объектный детектор, raw sensor fusion, safety monitor или сертифицированную систему управления автомобилем.",
                "Литературный пример BEVFusion не смешивается с собственным экспериментом. Он используется для объяснения влияния деградации сенсоров, тогда как собственные метрики относятся только к KITTI scenario table.",
            ],
        ),
        (
            "3.9 Воспроизводимость",
            [
                "Полный цикл повторяется командами подготовки данных, обучения, оценки, построения графиков и сборки документов. Все числовые результаты сохраняются в results, а итоговые документы берут метрики из этих файлов.",
                "Фиксированный seed 54 используется при разделении данных. Train, validation и test split сохранены в data/processed/kitti_split.json. Это позволяет повторить оценку и получить тот же набор сцен в каждом разбиении.",
                "Unit tests проверяют базовую логику ядра проекта. Отдельные проверки сверяют стиль текста и совпадение чисел в документах с results/metrics.json.",
            ],
        ),
        (
            "3.10 Сравнение с литературным примером BEVFusion",
            [
                "Публикации по BEVFusion показывают, что деградация отдельных сенсоров может заметно менять качество 3D object detection. Этот факт поддерживает общий мотив ВКР: reliability и uncertainty нельзя считать второстепенными признаками.",
                "Собственный эксперимент не обучает BEVFusion и не сравнивает скорость или точность с этой архитектурой. Здесь решается более узкая задача: оценить критичность сцены на основе реальных KITTI-аннотаций и воспроизводимой табличной модели.",
                "Такое разделение важно для корректности работы. Литературные данные помогают обосновать проблему, а собственные данные отвечают за экспериментальную часть.",
            ],
        ),
        (
            "3.11 Проверка целостности результатов",
            [
                "Числа в тексте, таблицах и презентации сверяются с results/metrics.json. Это защищает работу от расхождения между отчетом и последним запуском эксперимента.",
                "Проверяются primary_model, размеры split, precision, recall, F1, accuracy, ROC AUC, PR AUC и элементы матрицы ошибок. Если хотя бы одно значение отсутствует в финальных материалах, проверка должна завершиться ошибкой.",
                "Такая проверка особенно полезна на финальной стадии, когда текст, графики и презентация пересобираются несколько раз. Она снижает риск случайно оставить старую цифру в одном из файлов.",
            ],
        ),
        (
            "3.12 Перенос на raw sensor pipeline",
            [
                "Следующий инженерный шаг связан с обучением на изображениях, LiDAR и radar, а также с проверкой погодных условий и отказов сенсоров. Для этого потребуется другой датасет или расширенный сбор данных, потому что KITTI в текущем эксперименте не дает всех нужных каналов.",
                "Будущий raw sensor контур потребует отдельного выбора backend, проверки совместимости на Windows и независимого расчета метрик. В тексте ВКР этот этап не должен выглядеть как часть уже полученных результатов.",
                "Текущий прототип полезен как воспроизводимая основа. Он показывает, какие признаки и метрики стоит сохранить при переходе к более тяжелой модели, и где нужно усилить проверку перед промышленным применением.",
            ],
        ),
    ]
    sections = sections + [
        (
            "3.13 Подробный разбор FP и FN",
            [
                f"На test split primary model дала FP = {primary['fp']} и FN = {primary['fn']}. Эти числа важнее смотреть не только как элементы матрицы ошибок, но и как список конкретных сцен. Поэтому скрипт оценки сохраняет results/error_cases.csv, где каждая ошибка связана с признаками сцены и вероятностью модели.",
                "False positive чаще появляются в сценах с близким объектом, высоким risk_prior или заметной окклюзией. Модель в таких случаях реагирует осторожно и относит сцену к critical_scene, хотя derived target считает ее безопасной. Для прототипа тестирования ADAS такие ошибки приемлемее, чем пропуск критического случая, но их нельзя игнорировать.",
                "False negative опаснее. Они возникают, когда правило относит сцену к критическим, а модель не набирает выбранный порог. Обычно такие случаи связаны с пограничной дистанцией, умеренной окклюзией или комбинацией признаков, которые по отдельности не выглядят сильными. Именно такие сцены полезно возвращать в набор сложных примеров.",
                "Разбор ошибок помогает проверить смысл признаков. Если FN концентрируются около определенного диапазона расстояний, значит risk_prior или масштабирование дистанции нуждаются в уточнении. Если FP связаны с occlusion, можно отдельно проверить вклад camera_quality_proxy и uncertainty_proxy.",
                "В работе не выполняется ручное исправление меток после анализа ошибок. Это принципиально: error analysis объясняет поведение модели, но не меняет test split и не улучшает метрики задним числом. Такой порядок сохраняет честность результата.",
                "Практический итог анализа ошибок состоит в списке направлений для следующей версии. Нужны более богатые признаки динамики, проверка последовательностей кадров, raw image признаки, LiDAR depth consistency и отдельная оценка weather domain shift.",
            ],
        ),
        (
            "3.14 Репликация эксперимента",
            [
                "Репликация эксперимента начинается с одного источника данных: KITTI label_2. Если архив аннотаций лежит в data/external/training/label_2, скрипт prepare_data.py строит data/processed/kitti_scenarios.csv и kitti_split.json. Если структура каталога другая, подготовка должна завершиться понятной ошибкой, а не молча создать неполный набор.",
                "После подготовки train.py обучает модели и сохраняет веса, стандартизацию признаков, список признаков и пороги. evaluate.py рассчитывает итоговые метрики, predictions и error cases. make_figures.py и make_diagrams.py строят графики и схемы, которые затем используются в тексте и презентации.",
                "Воспроизводимость держится на том, что отчет не содержит вручную переписанных чисел. Значения precision, recall, F1, accuracy, ROC AUC, PR AUC, TP, FP, FN и TN берутся из results/metrics.json. Если эксперимент будет перезапущен с другими данными или seed, документы должны быть пересобраны.",
                "Unit tests проверяют базовую логику расчета риска и CLI. Они не заменяют экспериментальную оценку, но защищают ядро проекта от очевидных регрессий. Для учебного проекта этого достаточно, чтобы показать контролируемую инженерную сборку.",
                "Публичный репозиторий должен содержать только те файлы, которые помогают повторить расчет или понять результат. Проверочные отчеты, контактные листы рендера и промежуточные черновики полезны локально, но не должны становиться частью публичного состояния проекта.",
                "Такой порядок делает работу проверяемой. Можно открыть README, выполнить команды полного цикла и сопоставить полученные файлы с итоговыми материалами подачи.",
            ],
        ),
        (
            "3.15 Границы переноса на raw sensor pipeline",
            [
                "Переход к raw sensor pipeline меняет задачу. Вместо табличных признаков из аннотаций модель должна будет получать изображения, облака точек, radar-измерения, калибровки и временную синхронизацию. Это уже не простой logistic regression, а отдельный контур обучения и валидации.",
                "KITTI в текущем эксперименте не дает всех каналов, которые нужны для такой проверки. Поэтому перенос нельзя описывать как уже выполненный. Корректная формулировка: текущая работа проверяет scenario-level слой, а raw sensor pipeline остается следующим этапом.",
                "Для raw sensor версии потребуется другой набор данных или расширенная сборка. nuScenes подходит как кандидат, потому что содержит несколько модальностей и radar [7]. CARLA может помочь со стрессовыми сценами и управляемыми погодными условиями [8]. Эти источники нужно будет валидировать отдельно.",
                "Текущие метрики не зависят от будущего raw sensor контура. Такое разделение важно для честности: описание будущего pipeline не должно создавать впечатление, что в ВКР уже обучена тяжелая нейросетевая модель.",
                "Смысл текущего прототипа при переносе сохраняется в наборе проверяемых идей. Нужно оставить раздельный учет надежности наблюдения, неопределенности, априорного риска, настройку порога на validation и отдельный анализ FN. Меняется только уровень признаков и модель.",
                "Перед промышленным применением дополнительно нужны safety case, проверка отказов, независимые test sets, журналирование версий данных, контроль калибровки и юридическая оценка. ВКР не закрывает эти задачи, но формирует основу, от которой можно идти дальше без смешения уже проверенного и еще не проверенного.",
            ],
        ),
    ]
    extend_section(sections, "3.13", [
        "В файле `results/error_cases.csv` сохраняются не только тип ошибки и score, но и признаки сцены. Это позволяет смотреть на ошибку как на конкретный дорожный случай, а не как на безличную строку в confusion matrix.",
        "Часть FP возникает там, где сцена близка к опасной по геометрии, но derived target остается отрицательным. Для исследовательского прототипа такая осторожность объяснима, однако в промышленном контуре ее нужно ограничивать, чтобы система не выдавала слишком много предупреждений.",
        "FN требуют более строгого разбора. Если сцена содержит уязвимого участника или сильную окклюзию, но score не достигает порога, значит набор признаков не полностью передал риск этой ситуации.",
        "Разбор FP/FN также помогает отделить ошибку модели от ограничения целевой метки. Поскольку target является производным правилом, часть спорных случаев может находиться на границе самого правила, а не только на границе классификатора.",
        "В отчете поэтому не делается вывод, что модель “понимает опасность” в человеческом смысле. Корректнее сказать, что она воспроизводит заданное правило критичности и показывает, какие признаки помогают находить такие сцены.",
        "Следующий инженерный шаг состоит в разметке или сборе дополнительных сложных случаев, но такие данные должны добавляться до нового обучения и оцениваться на отдельном split, а не исправлять текущий test результат.",
    ])
    extend_section(sections, "3.14", [
        "Для повторения эксперимента важно сохранить порядок команд. Если сначала пересобрать документы, а затем изменить метрики, итоговые файлы могут стать несогласованными. Поэтому полный цикл начинается с подготовки данных и заканчивается сборкой материалов.",
        "Все промежуточные результаты имеют назначение: `models.json` хранит параметры моделей, `metrics.json` хранит итоговые числа, `predictions.json` позволяет перепроверить отдельные решения, а `error_cases.csv` связывает ошибки с признаками сцены.",
        "Графики строятся из тех же CSV/JSON, что и текст. Это снижает риск, что на рисунке останется старая метрика после повторного запуска обучения или оценки.",
        "Скрипт сборки документов не должен создавать публичные служебные чек-листы и отчеты внутренней проверки. Такие материалы полезны локально, но в репозитории они отвлекают от воспроизводимого проекта.",
        "Проверка репозитория после очистки включает поиск старых имен файлов, служебных отчетов и упоминаний промежуточных веток. Это не влияет на метрики, но влияет на восприятие работы как аккуратно подготовленного инженерного результата.",
        "Репликация также ограничена лицензией и доступностью KITTI. В репозиторий не добавляются тяжелые исходные архивы, но структура путей и ожидаемые выходные файлы описаны так, чтобы подготовку можно было повторить на собственной копии данных.",
    ])
    extend_section(sections, "3.15", [
        "Переход на raw sensor pipeline потребует другой архитектуры эксперимента. Табличная logistic regression удобна для интерпретации, но она не проверяет устойчивость нейросетевого детектора к шуму изображения, разрежению LiDAR или ошибке radar.",
        "В будущей версии нужно отделить perception model от risk assessment layer. Первый модуль будет извлекать объекты и confidence из сенсорных потоков, второй - оценивать надежность сцены и риск пропуска.",
        "Для обучения или инференса более тяжелых моделей нужно отдельно подбирать вычислительный backend и проверять его на Windows. Нельзя переносить текущие CPU-метрики на будущую нейросетевую архитектуру.",
        "Для проверки плохой погоды нужны данные, где погодные условия размечены или контролируются. KITTI в текущем виде не дает полноценной weather-разметки, поэтому любые выводы о тумане, снеге и дожде в экспериментальной части должны оставаться осторожными.",
        "При добавлении nuScenes, CARLA или собственного полигона потребуется заново определить split, метки, метрики и критерии исключения данных. Иначе результаты разных источников будут смешаны без общей базы сравнения.",
        "Именно поэтому текущая ВКР фиксирует проверяемый первый слой: real-data scenario table, derived target, интерпретируемая модель, анализ ошибок и ясные ограничения. Это не финальная ADAS-система, а основа для следующего этапа разработки.",
    ])
    extend_section(sections, "3.15", [
        "Отдельная сложность будущего raw sensor pipeline состоит в синхронизации каналов. Даже если камера и LiDAR по отдельности дают приемлемое качество, небольшая временная рассинхронизация может изменить положение объекта относительно траектории и тем самым изменить оценку риска.",
        "Для промышленного варианта потребуется контролировать калибровку. Табличная постановка работает с уже готовой геометрией из KITTI-аннотаций, а raw pipeline должен будет сам обеспечить согласование координат камеры, LiDAR и radar.",
        "Еще одна граница текущей работы связана с последовательностью кадров. В ВКР сцена рассматривается как отдельный снимок, тогда как автомобильная система обычно принимает решение с учетом динамики: скорости сближения, изменения occlusion и устойчивости детекции во времени.",
        "Из этого следует, что текущий прототип нельзя напрямую использовать как модуль принятия решения на дороге. Его корректная роль - фильтр и аналитический слой для поиска сложных сцен, которые нужно дополнительно проверять и включать в тестовые наборы.",
        "При переходе к нейросетевым моделям также меняется характер ошибок. Табличная модель ошибается из-за неполноты признаков и порога, а raw detector может ошибаться из-за доменного сдвига, артефактов изображения, слабой разметки или сбоя предварительной обработки.",
        "Поэтому сравнение будущей модели с текущей должно быть организовано как отдельный эксперимент. Нельзя просто добавить новую архитектуру и оставить прежний вывод без проверки на независимом наборе сцен.",
        "Сохранить из текущей работы стоит не конкретный классификатор, а дисциплину эксперимента: фиксированный split, явную целевую переменную, отдельный validation для порога, test only для финальной оценки, анализ FN и сохранение error cases.",
        "Такой перенос делает результат ВКР полезным для дальнейшей разработки: он не обещает готовое промышленное решение, но оставляет понятную инженерную карту, по которой можно двигаться к более тяжелой модели.",
        "Если на следующем этапе будет использоваться симулятор, синтетические сцены должны быть вынесены в отдельный набор и не смешиваться с KITTI-метриками. Иначе невозможно будет понять, что именно улучшило результат: новые данные, другой домен или изменение правила разметки.",
        "Если будет собран собственный набор, потребуется описать процедуру согласия, приватности и обезличивания. Для дорожных сцен это важно из-за возможного присутствия лиц, номеров автомобилей и геолокационных признаков.",
        "Отдельного внимания потребует валидация редких погодных условий. Для них мало обычного случайного split: нужны группы проверки, где туман, дождь, снег, ночь и блики рассматриваются как отдельные домены риска.",
        "В таком виде граница текущей ВКР становится ясной: здесь доказана воспроизводимая scenario-level методика на реальных KITTI-аннотациях, а не завершена вся цепочка восприятия автономного автомобиля.",
    ])
    write_clean_sections(doc, sections)
    table_caption(doc, 3, "Структура репозитория")
    add_table(doc, [
        ["Путь", "Назначение"],
        ["src/adas_scenarioguard", "ядро CLI и экспериментальные функции"],
        ["scripts", "подготовка данных, обучение, оценка, графики, документы"],
        ["data/processed", "таблица KITTI scenario table и split"],
        ["results", "модели, метрики, predictions, ошибки"],
        ["figures", "графики и диаграммы для ВКР"],
        ["docs", "воспроизводимость, анализ ошибок, вопросы защиты"],
    ], widths=[5, 11])
    table_caption(doc, 4, "Метрики моделей на test split KITTI")
    add_table(doc, [
        ["Метрика", "baseline", "proposed", "ablation"],
        ["Precision", baseline["precision"], primary["precision"], ablation["precision"]],
        ["Recall", baseline["recall"], primary["recall"], ablation["recall"]],
        ["F1", baseline["f1"], primary["f1"], ablation["f1"]],
        ["FNR", baseline["false_negative_rate"], primary["false_negative_rate"], ablation["false_negative_rate"]],
        ["ROC AUC", baseline["roc_auc"], primary["roc_auc"], ablation["roc_auc"]],
        ["PR AUC", baseline["pr_auc"], primary["pr_auc"], ablation["pr_auc"]],
    ], widths=[4, 4, 4, 4])
    doc.add_page_break()
    table_caption(doc, 5, "Матрица ошибок primary model")
    add_table(doc, [
        ["", "Predicted critical", "Predicted OK"],
        ["Actual critical", primary["tp"], primary["fn"]],
        ["Actual OK", primary["fp"], primary["tn"]],
    ], widths=[5, 5, 5])
    add_figure(doc, "component_diagram.png", 4, "Компонентная схема прототипа", 6.0)
    add_figure(doc, "deployment_diagram.png", 5, "Схема развертывания прототипа и будущего raw-sensor направления", 6.0)
    add_figure(doc, "confusion_matrix.png", 6, "Матрица ошибок primary model на KITTI test split", 5.4)
    add_figure(doc, "metrics_comparison.png", 7, "Сравнение baseline, proposed и ablation", 5.8)
    add_figure(doc, "roc_curve.png", 8, "ROC-кривая proposed model", 5.4)
    add_figure(doc, "precision_recall_curve.png", 9, "Precision-Recall кривая proposed model", 5.4)
    add_figure(doc, "error_by_condition.png", 10, "F1 по группам сцен", 5.8)
    add_figure(doc, "sensor_ablation.png", 11, "Ablation признаков на test split", 5.8)


def build_report() -> Path:
    doc = Document()
    setup_document(doc, page_numbers=False)
    m = metrics()
    summary = dataset_summary()
    add_title_page(doc)
    add_assignment_in_report_clean(doc)
    start_numbered_section(doc, 3)
    add_annotation(doc, m, summary, page_break=False)
    add_toc_clean(doc)
    add_terms(doc)
    add_intro(doc, m, summary)
    add_chapter_1_clean(doc)
    add_chapter_2_clean(doc)
    add_chapter_3_clean(doc, m, summary)
    add_conclusion(doc, m, summary)
    add_literature(doc)
    add_appendices(doc, m)
    set_core_properties(doc, "Выпускная квалификационная работа")
    FINAL.mkdir(exist_ok=True)
    path = FINAL / "Marianovskiy_VKR_ADAS_final.docx"
    doc.save(path)
    return path


def build_assignment() -> Path:
    doc = Document()
    setup_document(doc, page_numbers=False)
    add_assignment_in_report_clean(doc, page_break=False)
    set_core_properties(doc, "Задание на выполнение ВКР")
    FINAL.mkdir(exist_ok=True)
    path = FINAL / "Marianovskiy_zadanie_na_VKR_final.docx"
    doc.save(path)
    return path

def build_competency_index() -> Path:
    doc = Document()
    setup_document(doc)
    p(doc, "ПРЕДМЕТНЫЙ УКАЗАТЕЛЬ КОМПЕТЕНЦИЙ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, first_line=False)
    p(doc, f"ВКР: {THEME}")
    p(doc, f"Автор: {AUTHOR}, группа {GROUP}")
    add_table(doc, competency_rows(), widths=[3, 8, 5])
    p(doc, "Руководитель ВКР: ______________________", first_line=False)
    set_core_properties(doc, "Предметный указатель компетенций")
    FINAL.mkdir(exist_ok=True)
    path = FINAL / "Marianovskiy_competency_index_final.docx"
    doc.save(path)
    return path



def main() -> int:
    paths = [build_report(), build_assignment(), build_competency_index()]
    for path in paths:
        print(f"Saved: {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
