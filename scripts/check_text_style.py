"""Check final Russian text for banned phrases and mechanical wording."""
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document


ROOT = Path(__file__).resolve().parents[1]

BANNED_PHRASES = [
    "в современном мире",
    "на сегодняшний день",
    "бурное развитие технологий",
    "следует отметить",
    "немаловажным является",
    "таким образом можно сделать вывод",
    "данная работа посвящена",
    "комплексный подход позволяет",
    "актуальность обусловлена",
    "в теме соединяются три уровня",
    "фактор, связанный с",
    "аспект, связанный с",
    "условие, связанное с",
    "проверяющему важно увидеть",
    "комиссия должна видеть",
    "такой контур удобен для защиты",
    "внутренняя логика темы",
    "раздел выполняет контрольную функцию",
    "не расширяет эксперимент искусственно",
]

MECHANICAL_PATTERNS = [
    re.compile(r"\bперв(ый|ое)\s+фактор\b", re.IGNORECASE),
    re.compile(r"\bвтор(ой|ое)\s+фактор\b", re.IGNORECASE),
    re.compile(r"\bтрет(ий|ье)\s+фактор\b", re.IGNORECASE),
    re.compile(r"\bвывод\s*:\s*$", re.IGNORECASE),
]

BIBLIO_HEADERS = {
    "литература",
    "список использованных источников",
    "список литературы",
}


def docx_text(path: Path) -> list[str]:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if text:
                paragraphs.append(text)
    return paragraphs


def pptx_text(path: Path) -> list[str]:
    paragraphs: list[str] = []
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    with ZipFile(path) as archive:
        slide_names = sorted(name for name in archive.namelist() if name.startswith("ppt/slides/slide") and name.endswith(".xml"))
        for name in slide_names:
            root = ET.fromstring(archive.read(name))
            texts = [node.text for node in root.findall(".//a:t", ns) if node.text]
            if texts:
                paragraphs.append(" ".join(texts))
    return paragraphs


def plain_text(path: Path) -> list[str]:
    return [line.rstrip() for line in path.read_text(encoding="utf-8").splitlines()]


def extract_text(path: Path) -> list[str]:
    if path.suffix.lower() == ".docx":
        return docx_text(path)
    if path.suffix.lower() == ".pptx":
        return pptx_text(path)
    return plain_text(path)


def in_bibliography(paragraph: str, current: bool) -> bool:
    normalized = paragraph.strip().lower()
    if normalized in BIBLIO_HEADERS:
        return True
    return current


def check_path(path: Path) -> dict:
    issues = []
    bibliography = False
    paragraphs = extract_text(path)
    for index, paragraph in enumerate(paragraphs, start=1):
        bibliography = in_bibliography(paragraph, bibliography)
        lower = paragraph.lower()
        for phrase in BANNED_PHRASES:
            if phrase in lower:
                issues.append({"path": str(path), "paragraph": index, "type": "banned_phrase", "value": phrase})
        for pattern in MECHANICAL_PATTERNS:
            if pattern.search(paragraph):
                issues.append({"path": str(path), "paragraph": index, "type": "mechanical_pattern", "value": pattern.pattern})
        if not bibliography and ";" in paragraph:
            issues.append({"path": str(path), "paragraph": index, "type": "semicolon", "value": paragraph[:140]})
        if not bibliography and "--" in paragraph:
            issues.append({"path": str(path), "paragraph": index, "type": "double_dash", "value": paragraph[:140]})
    return {"path": str(path), "paragraphs": len(paragraphs), "issues": issues}


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("paths", nargs="*", help="DOCX, PPTX or text files to check")
    parser.add_argument("--json", dest="json_path", default=str(ROOT / "work" / "text_style_check.json"))
    args = parser.parse_args()

    if args.paths:
        paths = [Path(p) for p in args.paths]
    else:
        paths = [
            ROOT / "Marianovskiy_VKR_ADAS_final_rebuild.docx",
            ROOT / "Marianovskiy_VKR_ADAS_defense_final_rebuild.pptx",
            ROOT / "Marianovskiy_zadanie_na_VKR_final_rebuild.docx",
            ROOT / "Marianovskiy_competency_index_final_rebuild.docx",
        ]
    results = [check_path(path) for path in paths if path.exists()]
    all_issues = [issue for result in results for issue in result["issues"]]
    out_path = Path(args.json_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps({"checked": results, "issue_count": len(all_issues)}, ensure_ascii=False, indent=2), encoding="utf-8")
    if all_issues:
        print(json.dumps({"issue_count": len(all_issues), "issues": all_issues[:20]}, ensure_ascii=False, indent=2))
        return 1
    print(f"Style check passed for {len(results)} files")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
