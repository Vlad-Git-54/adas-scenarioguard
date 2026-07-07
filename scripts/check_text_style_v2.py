"""Check final clean VKR artifacts for template wording and local originality."""
from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document


ROOT = Path(__file__).resolve().parents[1]

TARGETS = [
    ROOT / "Marianovskiy_VKR_ADAS_final_clean.docx",
    ROOT / "Marianovskiy_zadanie_na_VKR_final_clean.docx",
    ROOT / "Marianovskiy_competency_index_final_clean.docx",
    ROOT / "Marianovskiy_VKR_ADAS_defense_final_clean.pptx",
]

BANNED_PHRASES = [
    "в современном мире",
    "на сегодняшний день",
    "бурное развитие технологий",
    "следует отметить",
    "немаловажным является",
    "таким образом можно сделать вывод",
    "данная работа посвящена",
    "комплексный подход позволяет",
    "актуальность обусловлена",
    "в теме ",
    "для темы ",
    "в разделе ",
    "для «",
    "раздел «",
    "тема «",
    "материал раздела",
    "содержание раздела",
    "фактор, связанный с",
    "аспект, связанный с",
    "условие, связанное с",
    "проверяющему важно увидеть",
    "проверяющий должен",
    "комиссия должна видеть",
    "удобен для защиты",
    "внутренняя логика темы",
    "не расширяет эксперимент искусственно",
    "Таблица 1 – Основные сведения задания".lower(),
]

MECHANICAL_PATTERNS = [
    re.compile(r"\bперв(ый|ое)\s+фактор\b", re.IGNORECASE),
    re.compile(r"\bвтор(ой|ое)\s+фактор\b", re.IGNORECASE),
    re.compile(r"\bтрет(ий|ье)\s+фактор\b", re.IGNORECASE),
    re.compile(r"\bвывод\s*:\s*$", re.IGNORECASE),
    re.compile(r"^(в\s+теме|для\s+темы|в\s+разделе|для\s+«|раздел\s+«|тема\s+«)", re.IGNORECASE),
]

BIBLIO_HEADERS = {
    "литература",
    "список использованных источников",
    "список литературы",
}


def docx_paragraphs(path: Path) -> list[str]:
    doc = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if text:
                paragraphs.append(text)
    return paragraphs


def pptx_paragraphs(path: Path) -> list[str]:
    ns = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    paragraphs: list[str] = []
    with ZipFile(path) as archive:
        names = [
            name
            for name in archive.namelist()
            if (name.startswith("ppt/slides/slide") or name.startswith("ppt/notesSlides/notesSlide"))
            and name.endswith(".xml")
        ]
        for name in sorted(names):
            root = ET.fromstring(archive.read(name))
            texts = [node.text for node in root.findall(".//a:t", ns) if node.text]
            if texts:
                paragraphs.append(" ".join(texts))
    return paragraphs


def extract(path: Path) -> list[str]:
    if path.suffix.lower() == ".docx":
        return docx_paragraphs(path)
    if path.suffix.lower() == ".pptx":
        return pptx_paragraphs(path)
    return path.read_text(encoding="utf-8").splitlines()


def is_bibliography_header(text: str) -> bool:
    return text.strip().lower() in BIBLIO_HEADERS


def check_target(path: Path) -> tuple[list[dict], list[str]]:
    issues: list[dict] = []
    paragraphs = extract(path) if path.exists() else []
    if not path.exists():
        return ([{"path": str(path), "paragraph": 0, "type": "missing_file", "value": path.name}], paragraphs)

    bibliography = False
    for index, paragraph in enumerate(paragraphs, start=1):
        if is_bibliography_header(paragraph):
            bibliography = True
        lower = paragraph.lower()
        if "codex" in lower:
            issues.append({"path": str(path), "paragraph": index, "type": "forbidden_name", "value": "codex"})
        if "--" in paragraph:
            issues.append({"path": str(path), "paragraph": index, "type": "double_dash", "value": "--"})
        if ";" in paragraph and not bibliography:
            issues.append({"path": str(path), "paragraph": index, "type": "semicolon", "value": ";"})
        for phrase in BANNED_PHRASES:
            if phrase in lower:
                issues.append({"path": str(path), "paragraph": index, "type": "banned_phrase", "value": phrase})
        for pattern in MECHANICAL_PATTERNS:
            if pattern.search(paragraph):
                issues.append({"path": str(path), "paragraph": index, "type": "mechanical_pattern", "value": pattern.pattern})
    return issues, paragraphs


def originality_estimate(paragraphs: list[str]) -> dict:
    words = re.findall(r"[A-Za-zА-Яа-яЁё0-9]+", " ".join(paragraphs).lower())
    if len(words) < 8:
        return {"word_count": len(words), "local_originality_percent": 100.0, "repeated_ngram_share": 0.0}
    ngrams = [tuple(words[i : i + 8]) for i in range(len(words) - 7)]
    counts = Counter(ngrams)
    repeated = sum(count - 1 for count in counts.values() if count > 1)
    share = repeated / len(ngrams)
    return {
        "word_count": len(words),
        "local_originality_percent": round((1.0 - share) * 100, 2),
        "repeated_ngram_share": round(share, 4),
    }


def main() -> int:
    all_issues: list[dict] = []
    all_paragraphs: list[str] = []
    checked: list[str] = []
    for target in TARGETS:
        issues, paragraphs = check_target(target)
        all_issues.extend(issues)
        all_paragraphs.extend(paragraphs)
        checked.append(target.name)

    originality = originality_estimate(all_paragraphs)
    status = "passed" if not all_issues and originality["local_originality_percent"] >= 70 else "failed"
    lines = [
        "# STYLE CHECK V2",
        "",
        "Checked files:",
        *[f"- `{name}`" for name in checked],
        "",
        f"Status: {status}",
        f"Local originality estimate: {originality['local_originality_percent']}%",
        f"Word count checked: {originality['word_count']}",
        f"Repeated 8-gram share: {originality['repeated_ngram_share']}",
        "",
        "This is a local lexical check, not an official university originality report.",
        "",
    ]
    if all_issues:
        lines.extend(["## Issues", ""])
        for issue in all_issues[:200]:
            lines.append(f"- `{Path(issue['path']).name}`, paragraph {issue['paragraph']}: {issue['type']} = `{issue['value']}`")
    else:
        lines.append("No banned wording, template openings, semicolons or double dashes were found in checked final artifacts.")

    (ROOT / "STYLE_CHECK_V2.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    (ROOT / "STYLE_CHECK_V2.json").write_text(
        json.dumps({"status": status, "issues": all_issues, "originality": originality}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(f"Style check V2 {status}")
    return 0 if status == "passed" else 1


if __name__ == "__main__":
    raise SystemExit(main())
